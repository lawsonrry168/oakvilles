# -*- coding: utf-8 -*-
"""Generate merged quotation: website + marketing summary + value comparison."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os
import shutil

BASE = r"c:\Users\Steriod\Desktop\oakvilles"
OUT = os.path.join(BASE, "Oakville-Combined-Quotation.docx")
OUT_PRES = os.path.join(BASE, "presentations", "Oakville-Combined-Quotation.docx")

PRICE_FACTOR = 0.7
WEBSITE_SUBTOTAL = int(49_000 * PRICE_FACTOR)
IMAGE_COUNT = 50
IMAGE_UNIT = int(380 * PRICE_FACTOR)
IMAGE_INTEGRATION = int(3_800 * PRICE_FACTOR)
IMAGE_SUBTOTAL = IMAGE_COUNT * IMAGE_UNIT + IMAGE_INTEGRATION
BUNDLE_DISCOUNT = int(3_800 * PRICE_FACTOR)
LIST_TOTAL = WEBSITE_SUBTOTAL + IMAGE_SUBTOTAL - BUNDLE_DISCOUNT
REFERRAL_DISCOUNT = int(LIST_TOTAL * 0.5)
REFERRAL_TOTAL = LIST_TOTAL - REFERRAL_DISCOUNT

MONTHLY_RETAINER = 8_000
MONTHLY_RETAINER_PREV = 10_000
MIN_CONTRACT_MONTHS = 3
HOURLY_DEV = int(650 * PRICE_FACTOR)
HOURLY_IMAGE = int(380 * PRICE_FACTOR)

META_FB_ALLOC = 1_000
META_IG_ALLOC = 1_000
GOOGLE_ALLOC = 3_000
# Remaining 3,000 of 8,000 = content + strategy + maintenance (content explicitly included)

today = datetime.date(2026, 7, 1)
valid_until = today + datetime.timedelta(days=30)


def fmt(n: int) -> str:
    return f"{n:,}"


def payment_split(total: int) -> tuple[int, int, int]:
    deposit = int(total * 0.4)
    mid = int(total * 0.4)
    final = total - deposit - mid
    return deposit, mid, final


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def style_run(run, size=10.5, bold=False, color=None):
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 14, True, RGBColor(0x2A, 0x46, 0x3C))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 11.5, True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def para(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size, bold)
    p.paragraph_format.space_after = Pt(4)


def bullets(doc, items, size=10):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            style_run(r, size)


def table(doc, headers, rows, col_widths=None, header_fill="D5E8F0"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        set_cell_shading(t.rows[0].cells[i], header_fill)
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                style_run(r, 9, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    style_run(r, 9)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def qa_block(doc, items):
    for q, a in items:
        p = doc.add_paragraph()
        rq = p.add_run(f"Q：{q}\n")
        style_run(rq, 10, True)
        ra = p.add_run(f"A：{a}")
        style_run(ra, 10)
        p.paragraph_format.space_after = Pt(8)


doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

# ── Cover ──
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = title_p.add_run("綜 合 報 價 單\n")
style_run(r1, 22, True, RGBColor(0x2A, 0x46, 0x3C))
r2 = title_p.add_run("網站重建 · 視覺資產 · 月度網絡行銷")
style_run(r2, 13)

doc.add_paragraph()
table(doc, ["欄位", "內容"], [
    ("報價編號", "SV-WEB-2026-0701"),
    ("報價日期", f"{today.year} 年 {today.month} 月 {today.day} 日"),
    ("有效期限", f"30 日（至 {valid_until.year} 年 {valid_until.month} 月 {valid_until.day} 日）"),
    ("幣別", "港元（HKD）"),
    ("客戶", "頤安本草 · 伍厚臻中醫師（Oakville Wellness）"),
    ("網域", "oakvilles.com ｜ oakvilles.vercel.app（已部署）"),
    ("承辦方", "[承辦方名稱]　｜　聯絡：[email] · [電話]"),
], [4, 12])

# ── Part A: Value comparison ──
h1(doc, "甲、價值前後比較（為何值得投資）")
para(
    doc,
    "以下對照舊站、行業標竿（養康）及新站現況，說明一次性重建與月度行銷之商業價值。",
)
h2(doc, "核心指標")
table(doc, ["指標", "舊站（重建前）", "新站（現況）", "價值提升"], [
    ("頁面規模", "實質 1 頁", "62 頁（中+英）", "可承接長尾搜尋與廣告著陸"),
    ("SEO 技術", "無 sitemap／Schema", "sitemap·JSON-LD·OG·hreflang", "搜尋與分享基礎完備"),
    ("轉化動線", "僅 WA 浮窗", "2 步 funnel·計算器·sticky CTA", "降低預約摩擦、可追蹤"),
    ("信任元件", "缺評價·環境·FAQ", "Google 5.0·gallery·明碼收費", "提升中環高端客群信心"),
    ("量度能力", "無 GA4", "6 個 dataLayer 事件就緒", "廣告 ROI 可量度"),
    ("品牌定位", "個人名義·訊息分散", "頤安本草·合規克制", "與促銷型對手區隔"),
], [2.5, 3.5, 4.5, 5.5])

h2(doc, "舊站 → 新站：8 項具體改善")
table(doc, ["類別", "舊站", "新站"], [
    ("頁面", "1 頁為主", "62 頁（含 /en/ 雙語）"),
    ("預約", "WA 連結", "2 步 funnel + WhatsApp 統一模板"),
    ("SEO", "基本 title", "8 症狀頁 + central-hk + 全站 meta"),
    ("收費", "未展示", "明碼價 + 即時計算器"),
    ("部署", "舊 hosting", "GitHub → Vercel 自動建置"),
    ("內容", "2 專科卡片", "4 專科 + 症狀頁 + Blog + FAQ"),
    ("合規", "有效能承諾", "禁止誇大表述"),
    ("手機", "基本 RWD", "sticky CTA · WA 浮動避障"),
], [2, 5.5, 8])

h2(doc, "相對養康：借鑑 vs 不照搬")
table(doc, ["借鑑（架構）", "不照搬（調性）", "頤安差異化"], [
    ("症狀著陸頁 SEO", "美容塑形促銷主入口", "信任 + 合規 + 高端體驗"),
    ("WhatsApp 全站轉化", "首診免診金", "中環專業定位"),
    ("FAQ + Blog 內容厚度", "論壇軟文佈局", "循本溯源品牌語氣"),
], [5, 5, 6])

# ── Part B: Quote summary ──
h1(doc, "乙、報價摘要")
table(doc, ["項目", "報價 HKD", "轉介優惠 HKD", "備註"], [
    ("一次性：網站 + 50 張圖", fmt(LIST_TOTAL), fmt(REFERRAL_TOTAL), "轉介客戶五折"),
    ("月度：行銷 + 維護", f"{fmt(MONTHLY_RETAINER)}/月", f"{fmt(MONTHLY_RETAINER)}/月",
     f"原 HKD {fmt(MONTHLY_RETAINER_PREV)}/月；最少 {MIN_CONTRACT_MONTHS} 個月"),
    ("平台廣告投放費", "客戶直付", "客戶直付", "代操含於月費；建議預算見下文"),
], [4.5, 2.8, 2.8, 5.9])

bullets(doc, [
    "新站現況：62 頁雙語、WhatsApp 6734 9532 全站統一、Vercel Production 已上線",
    "轉化 KPI：whatsapp_click（全站 data-cta-id 可追蹤）",
    "月度內容：8 則社交（IG 4 + FB 4）+ 2 篇官網文章 — 已含於月費 HKD 8,000",
])

# ── Part C: One-time ──
h1(doc, "丙、一次性項目（網站 + 視覺）")
h2(doc, "網站開發範圍（已部署基礎版）")
table(doc, ["模組", "交付內容"], [
    ("資訊架構", "62 頁靜態站（繁中 + /en/）；專科、症狀、Blog、FAQ、流程收費"),
    ("轉化動線", "2 步 WhatsApp 預約、診金計算器、sticky CTA"),
    ("SEO 與部署", "sitemap、OG 1200×630、Schema、Vercel、DNS 指引"),
    ("視覺資產", f"{IMAGE_COUNT} 張品牌圖 + 全站 WebP 接入（含 2 輪修訂）"),
], [3.5, 12])

h2(doc, "一次性費用")
table(doc, ["類別", "小計 HKD"], [
    ("網站開發", fmt(WEBSITE_SUBTOTAL)),
    (f"視覺資產（{IMAGE_COUNT} 張 + 接入）", fmt(IMAGE_SUBTOTAL)),
    ("打包優惠", f"−{fmt(BUNDLE_DISCOUNT)}"),
    ("報價總額", fmt(LIST_TOTAL)),
    ("轉介優惠（五折）", f"−{fmt(REFERRAL_DISCOUNT)}"),
    ("轉介價", fmt(REFERRAL_TOTAL)),
], [10, 6])

dep_std, mid_std, fin_std = payment_split(LIST_TOTAL)
dep_ref, mid_ref, fin_ref = payment_split(REFERRAL_TOTAL)
h2(doc, "付款里程碑（一次性）")
table(doc, ["階段", "比例", "報價 HKD", "轉介 HKD", "觸發"], [
    ("訂金", "40%", fmt(dep_std), fmt(dep_ref), "合約簽署"),
    ("中期", "40%", fmt(mid_std), fmt(mid_ref), "Staging 驗收 + 首批 20 張圖"),
    ("尾款", "20%", fmt(fin_std), fmt(fin_ref), "上線 + 50 張圖全站接入"),
], [2, 1.5, 2.5, 2.5, 7])

# ── Part D: Monthly marketing ──
h1(doc, "丁、月度網絡行銷及維護")
para(
    doc,
    f"月費 HKD {fmt(MONTHLY_RETAINER)}（原 HKD {fmt(MONTHLY_RETAINER_PREV)}），"
    f"最少合約 {MIN_CONTRACT_MONTHS} 個月，按月預付。"
    "內容製作（8 則社交 + 2 篇官網文章）已含於月費，不再另收 HKD 3,000–6,000。",
    True,
)

h2(doc, "月費配置明細（服務費 HKD 8,000）")
table(doc, ["項目", "月費配置 HKD", "說明"], [
    ("Meta Facebook 代操", fmt(META_FB_ALLOC), "活動設定、受眾、素材測試、Pixel 追蹤、月報"),
    ("Meta Instagram 代操", fmt(META_IG_ALLOC), "Reels/Stories、IG 專屬受眾、與 FB 協同優化"),
    ("Google 搜尋代操", fmt(GOOGLE_ALLOC), "品牌+中環+症狀 Search；著陸 /conditions/ 及 central-hk"),
    ("內容製作", "月費包含", "每月 4 IG + 4 FB post + 2 篇官網 Blog（800–1200 字）"),
    ("策略·維護·會議", "月費包含", "GBP 建議、GA4 月報、Vercel 小修、每月 60 分鐘檢討"),
    ("合計", fmt(MONTHLY_RETAINER), "不含平台廣告投放費（客戶直付 Meta / Google）"),
], [4.5, 2.5, 9])

h2(doc, "每月固定產出")
table(doc, ["渠道", "頻率", "產出"], [
    ("Instagram", "每週 1 篇", "4 post（症狀科普·診所信任·季節養生·預約 CTA）"),
    ("Facebook", "每週 1 篇", "4 post（可共用 IG 素材，依平台微調）"),
    ("官網 Blog", "每兩週 1 篇", "2 篇文章（SEO 長尾 + 內部連結症狀頁）"),
], [3, 3, 10])

h2(doc, "建議廣告投放預算（客戶直付平台，參考）")
para(doc, "以下為試跑期建議 spend，與月費代操服務分開；可依 whatsapp_click 成本調整。", size=10)
table(doc, ["平台", "建議月投放 HKD", "備註"], [
    ("Meta Facebook", fmt(META_FB_ALLOC), "症狀受眾 + 再行銷"),
    ("Meta Instagram", fmt(META_IG_ALLOC), "Reels/圖文導流症狀頁"),
    ("Google 搜尋", fmt(GOOGLE_ALLOC), "高意向關鍵字"),
    ("內容製作", "—", "已含於月費 HKD 8,000"),
    ("客戶每月合計（服務+廣告）", fmt(MONTHLY_RETAINER + META_FB_ALLOC + META_IG_ALLOC + GOOGLE_ALLOC),
     f"服務 {fmt(MONTHLY_RETAINER)} + 廣告 {fmt(META_FB_ALLOC + META_IG_ALLOC + GOOGLE_ALLOC)}"),
], [4.5, 3, 8.5])

h2(doc, "價格調整前後對照")
table(doc, ["項目", "調整前", "調整後"], [
    ("月度服務費", f"HKD {fmt(MONTHLY_RETAINER_PREV)}/月", f"HKD {fmt(MONTHLY_RETAINER)}/月"),
    ("內容製作（8社交+2文）", "另計 HKD 3,000–6,000", "月費包含"),
    ("Meta 代操", "合併列「Meta 5,000–8,000」", f"FB {fmt(META_FB_ALLOC)} + IG {fmt(META_IG_ALLOC)}"),
    ("Google 代操", "5,000–8,000（參考）", f"HKD {fmt(GOOGLE_ALLOC)}（配置）"),
    ("建議客戶月總支出", "約 13,000–22,000", f"約 {fmt(MONTHLY_RETAINER + META_FB_ALLOC + META_IG_ALLOC + GOOGLE_ALLOC)}（服務+試跑廣告）"),
], [5, 5, 6])

# ── Part E: Marketing execution summary (concise) ──
h1(doc, "戊、行銷執行摘要")
para(
    doc,
    "一句話：用社交內容 + 搜尋廣告把人帶到官網，最後統一經 WhatsApp 6734 9532 預約。",
    True,
)

table(doc, ["重點", "說明"], [
    ("做什麼", "每月 4 則 IG、4 則 FB、2 篇官網文章；代操 Meta FB／IG、Google 搜尋"),
    ("給誰看", "中環、金鐘、半山上班族；另有英文站 /en/ 承接外籍客群"),
    ("賣什麼", "濕疹、暗瘡、失眠、備孕、頸肩痛等症狀專頁"),
    ("怎樣算成功", "用戶點擊 WhatsApp（whatsapp_click）— 其餘數據輔助優化"),
    ("不做什麼", "免診金、論壇軟文、豐胸減肥促銷（與高端定位不符）"),
], [2.8, 13.2])

h2(doc, "每月四週（固定節奏）")
bullets(doc, [
    "第 1 週：症狀科普貼文",
    "第 2 週：診所信任貼文 + 發第 1 篇 Blog",
    "第 3 週：季節養生貼文 + 加碼表現最佳廣告",
    "第 4 週：預約教學貼文 + 發第 2 篇 Blog + 月度數據檢討",
    "每則內容結尾 CTA：WhatsApp 6734 9532",
], size=10.5)

h2(doc, "廣告點擊後去哪一頁")
table(doc, ["想推廣", "打開頁面"], [
    ("濕疹／暗瘡", "官網症狀頁"),
    ("備孕／婦科", "官網備孕頁"),
    ("中環找中醫", "中環專頁 central-hk"),
    ("外籍客人", "英文首頁或皮膚專科頁"),
], [4, 12])

h2(doc, "上線後三步走")
bullets(doc, [
    "第 1–2 週：接好 GA4、廣告追蹤、Google 商家、網域切換",
    "第 3–6 週：Meta + Google 試跑，看 WhatsApp 點擊成本",
    "第 2–3 月：保留有效廣告，持續每月內容更新",
], size=10.5)

# ── Part F: Terms (abbreviated) ──
h1(doc, "己、條款與 FAQ（摘要）")
bullets(doc, [
    f"月度合約最少 {MIN_CONTRACT_MONTHS} 個月；廣告費客戶直付平台，代操含於月費 HKD {fmt(MONTHLY_RETAINER)}",
    "不保證預約量或 ROAS；義務為約定產出 + 代操 + 數據報告",
    "文案合規由客戶最終確認；不含 KOL 費、醫師現場攝影、法律意見",
    f"超範圍開發按 HKD {HOURLY_DEV}/hr 另計；本報價 30 日有效",
    "帳戶（IG/FB/Google）應登記於客戶名下；合約終止須交還管理權限",
])

h2(doc, "不包含（另計）")
table(doc, ["項目", "說明"], [
    ("Meta / Google 廣告投放費", "客戶直付；建議 FB 1K + IG 1K + Google 3K/月試跑"),
    ("KOL 合作費", "另列專項預算"),
    ("Vercel Pro、網域年費", "各約 USD 20+/月、HKD 200–400/年"),
], [6, 10])

# ── Signature ──
h1(doc, "庚、確認簽署")
sig = doc.add_table(rows=5, cols=2)
sig.style = "Table Grid"
sig.rows[0].cells[0].text = "客戶 Client"
sig.rows[0].cells[1].text = "承辦方 Provider"
sig.rows[1].cells[0].text = "頤安本草 · 伍厚臻中醫師"
sig.rows[1].cells[1].text = "[承辦方名稱]"
sig.rows[2].cells[0].text = "授權代表\n\n_________________________"
sig.rows[2].cells[1].text = "授權代表\n\n_________________________"
sig.rows[3].cells[0].text = "日期\n\n_________________________"
sig.rows[3].cells[1].text = "日期\n\n_________________________"
sig.rows[4].cells[0].text = "簽署\n\n_________________________"
sig.rows[4].cells[1].text = "簽署\n\n_________________________"

doc.add_paragraph()
para(
    doc,
    f"本文件合併《Oakville-Website-Quotation》與《Oakville-Marketing-Summary》之精華，"
    f"含價值前後比較。一次性 HKD {fmt(LIST_TOTAL)}（轉介 {fmt(REFERRAL_TOTAL)}）；"
    f"月度 HKD {fmt(MONTHLY_RETAINER)}/月。",
    size=8,
)

doc.save(OUT)
os.makedirs(os.path.dirname(OUT_PRES), exist_ok=True)
shutil.copy2(OUT, OUT_PRES)
print(f"Saved: {OUT}")
print(f"Copied: {OUT_PRES}")
