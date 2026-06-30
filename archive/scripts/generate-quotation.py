# -*- coding: utf-8 -*-
"""Generate restructured quotation docx for 頤安本草."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os
import shutil

OUTPUT = os.path.join(r"c:\Users\Steriod\Desktop\oakvilles", "Oakville-Website-Quotation.docx")
OUTPUT_DESKTOP = r"c:\Users\Steriod\Desktop\Oakville-Website-Quotation.docx"
OUTPUT_PRESENTATIONS = os.path.join(
    r"c:\Users\Steriod\Desktop\oakvilles\presentations", "Oakville-Website-Quotation.docx"
)

PRICE_FACTOR = 0.7  # 全項減價 30%

WEBSITE_SUBTOTAL = int(49_000 * PRICE_FACTOR)
IMAGE_COUNT = 50
IMAGE_UNIT = int(380 * PRICE_FACTOR)
IMAGE_INTEGRATION = int(3_800 * PRICE_FACTOR)
IMAGE_SUBTOTAL = IMAGE_COUNT * IMAGE_UNIT + IMAGE_INTEGRATION
BUNDLE_DISCOUNT = int(3_800 * PRICE_FACTOR)
LIST_TOTAL = WEBSITE_SUBTOTAL + IMAGE_SUBTOTAL - BUNDLE_DISCOUNT

REFERRAL_DISCOUNT = int(LIST_TOTAL * 0.5)
REFERRAL_TOTAL = LIST_TOTAL - REFERRAL_DISCOUNT

MONTHLY_RETAINER = 10_000  # 月度服務不適用報價折扣
MIN_CONTRACT_MONTHS = 3
HOURLY_DEV = int(650 * PRICE_FACTOR)
HOURLY_IMAGE = int(380 * PRICE_FACTOR)

today = datetime.date(2026, 6, 23)
valid_until = today + datetime.timedelta(days=30)


def fmt(n: int) -> str:
    return f"{n:,}"


def payment_split(total: int) -> tuple[int, int, int]:
    deposit = int(total * 0.4)
    mid = int(total * 0.4)
    final = total - deposit - mid
    return deposit, mid, final


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def style_run(run, size=10.5, bold=False, color=None):
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 14, True, RGBColor(0x2A, 0x46, 0x3C))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 11.5, True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def para(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size, bold)
    p.paragraph_format.space_after = Pt(4)


def bullets(doc, items, size=10):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            style_run(r, size)


def table(doc, headers, rows, col_widths=None, header_fill="D5E8F0"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        set_cell_shading(t.rows[0].cells[i], header_fill)
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                style_run(r, 9, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    style_run(r, 9)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def qa_block(doc, items):
    """Q&A pairs: list of (question, answer)"""
    for q, a in items:
        p = doc.add_paragraph()
        rq = p.add_run(f"Q：{q}\n")
        style_run(rq, 10, True)
        ra = p.add_run(f"A：{a}")
        style_run(ra, 10)
        p.paragraph_format.space_after = Pt(8)


doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

# ── 封面 ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = title.add_run("報 價 單\n")
style_run(r1, 22, True, RGBColor(0x2A, 0x46, 0x3C))
r2 = title.add_run("QUOTATION")
style_run(r2, 14)

doc.add_paragraph()
table(doc, ["欄位", "內容"], [
    ("報價編號", "SV-WEB-2026-0623"),
    ("報價日期", f"{today.year} 年 {today.month} 月 {today.day} 日"),
    ("有效期限", f"30 日（至 {valid_until.year} 年 {valid_until.month} 月 {valid_until.day} 日）"),
    ("幣別", "港元（HKD）；未含第三方平台及網域年費"),
    ("客戶", "頤安本草 · 伍厚臻中醫師（Oakville Wellness）"),
    ("專案", "官方診所網站重建、視覺資產及月度網絡行銷"),
    ("網域", "oakvilles.com"),
    ("承辦方", "[承辦方名稱]　｜　聯絡：[email] · [電話]"),
], [4, 12])

# ── 一、報價摘要 ──
h1(doc, "一、報價摘要")
para(
    doc,
    "本報價分兩部分：① 一次性項目（網站 + 50 張視覺資產）；② 月度網絡行銷及維護（另簽約、按月計費）。"
    "以下為重點速覽，細節見後文。",
)
table(doc, ["項目", "報價 HKD", "轉介優惠 HKD", "備註"], [
    ("一次性：網站 + 50 張圖", fmt(LIST_TOTAL), fmt(REFERRAL_TOTAL), "轉介客戶五折；簽約時確認即可"),
    ("月度：行銷 + 維護", f"{fmt(MONTHLY_RETAINER)}/月", f"{fmt(MONTHLY_RETAINER)}/月", f"最少 {MIN_CONTRACT_MONTHS} 個月；不含廣告費"),
    ("Meta / Google 廣告費", "—", "—", "客戶直接支付平台，代操含於月費"),
], [4.5, 2.8, 2.8, 5.9])

bullets(doc, [
    "新站現況：62 頁（繁中 + 英文 /en/）、WhatsApp 預約 funnel、GA4 事件就緒、Vercel 已部署",
    "轉化 KPI：WhatsApp 6734 9532（全站統一）",
    "本報價按已實作範圍及交付清單報價；正式合約以雙方確認之工作說明書（SOW）為準",
])

# ── 二、一次性項目 ──
h1(doc, "二、一次性項目（網站 + 視覺）")

h2(doc, "2.1 網站開發範圍")
table(doc, ["模組", "交付內容"], [
    ("資訊架構", "62 頁靜態站（繁中 + /en/）；專科、症狀、Blog、FAQ、流程收費等"),
    ("轉化動線", "2 步 WhatsApp 預約、診金計算器、sticky CTA、浮動 WA 按鈕"),
    ("互動功能", "全站搜尋、dataLayer 6 事件、Schema.org 結構化資料"),
    ("SEO 與部署", "sitemap、OG meta、hreflang、Vercel Production、DNS 指引"),
    ("驗收", "跨裝置 UAT、預約動線測試、上線確認清單"),
], [3.5, 12])

h2(doc, "2.2 視覺資產（50 張）")
table(doc, ["類別", "數量", "說明"], [
    ("診所實景", "6", "候診區、診症室、針灸區等；替換佔位圖"),
    ("醫師肖像", "3", "以客戶提供 doctor.jpg 優化（非憑空生成人臉）"),
    ("專欄 / 社交", "10", "Blog 封面、IG 方圖等"),
    ("專科 / 症狀 Hero", "14", "專科頁 + 症狀頁主視覺"),
    ("流程 / OG / 雜項", "17", "流程步驟、FAQ 配圖、各頁 OG 分享圖"),
    ("後製接入", "—", "WebP + alt + 全站 HTML 替換；含 2 輪修訂"),
], [3.5, 1.5, 11])

# ── 三、月度網絡行銷（依現行行銷方案）──
h1(doc, "三、月度網絡行銷及維護")
para(
    doc,
    f"月費 HKD {fmt(MONTHLY_RETAINER)}，最少合約 {MIN_CONTRACT_MONTHS} 個月，按月預付。"
    "以下按《Oakville-Digital-Marketing-Plan》執行；廣告投放費由客戶直接支付 Meta / Google，不含於月費。",
    True,
)

h2(doc, "3.1 每月固定產出")
table(doc, ["渠道", "頻率", "每月產出", "說明"], [
    ("Instagram", "每週 1 篇", "4 post", "症狀科普、診所信任、季節養生、預約 CTA 四大支柱輪替"),
    ("Facebook", "每週 1 篇", "4 post", "可與 IG 共用素材，依平台微調文案與 CTA"),
    ("官網養生專欄", "每兩週 1 篇", "2 篇文章", "800–1200 字；SEO 長尾 + 內部連結症狀頁"),
], [2.5, 2, 2, 9])

h2(doc, "3.2 廣告代操（不含廣告費）")
table(doc, ["平台", "服務內容", "著陸頁方向"], [
    ("Meta（FB + IG）", "活動設定、受眾、素材測試、Pixel 追蹤、月報", "症狀頁（濕疹/暗瘡/備孕/失眠）+ 再行銷"),
    ("Google Ads", "Search（品牌/中環/症狀）+ PMax 轉化優化、月報", "/conditions/ 及 central-hk"),
], [3, 7, 6])

h2(doc, "3.3 策略、量度與技術維護")
table(doc, ["#", "類別", "每月服務"], [
    ("M1", "本地 SEO", "Google Business Profile 優化建議、Search Console 監測"),
    ("M2", "轉化分析", "GA4 / GTM 維護、whatsapp_click 月報、漏斗優化建議"),
    ("M3", "網站維護", "Vercel 監控、SSL/DNS 檢查、小修版式（≤4 張圖替換）"),
    ("M4", "策略會議", "每月 1 次 60 分鐘線上檢討（數據 + 下月計劃）"),
], [1, 3, 12])

h2(doc, "3.4 合規與不包含")
bullets(doc, [
    "文案合規：不採免診金、論壇軟文、豐胸減肥等與品牌定位不符手段；醫療表述由客戶最終確認",
    "不含：Meta / Google 廣告投放費、KOL 費、醫師現場攝影、法律合規意見",
    f"超範圍：額外頁面、大量改版、超 2 篇 Blog → 按 HKD {HOURLY_DEV}/hr 或另議",
])

# ── 四、費用明細 ──
h1(doc, "四、費用明細")

h2(doc, "4.1 一次性項目")
table(doc, ["類別", "小計 HKD"], [
    ("網站開發（策略、前端、互動、SEO、部署）", fmt(WEBSITE_SUBTOTAL)),
    (f"視覺資產（{IMAGE_COUNT} 張 + 全站接入）", fmt(IMAGE_SUBTOTAL)),
    ("打包優惠", f"−{fmt(BUNDLE_DISCOUNT)}"),
    ("報價總額", fmt(LIST_TOTAL)),
    ("轉介優惠（五折）", f"−{fmt(REFERRAL_DISCOUNT)}"),
    ("轉介價", fmt(REFERRAL_TOTAL)),
], [10, 6])

h2(doc, "4.2 月度服務")
table(doc, ["項目", "費用"], [
    ("月度行銷 + 維護（M1–M4 + 3.1–3.2）", f"HKD {fmt(MONTHLY_RETAINER)}/月"),
    ("最少合約期", f"{MIN_CONTRACT_MONTHS} 個月"),
    ("超時開發 / 設計", f"HKD {HOURLY_DEV}/hr 另計"),
    ("額外圖片重製（超 2 輪）", f"HKD {HOURLY_IMAGE}/張 另計"),
], [10, 6])

# ── 五、付款方式 ──
h1(doc, "五、付款方式")

dep_std, mid_std, fin_std = payment_split(LIST_TOTAL)
dep_ref, mid_ref, fin_ref = payment_split(REFERRAL_TOTAL)

h2(doc, "5.1 一次性項目")
table(doc, ["階段", "比例", "報價 HKD", "轉介優惠 HKD", "觸發條件"], [
    ("訂金", "40%", fmt(dep_std), fmt(dep_ref), "合約簽署"),
    ("中期", "40%", fmt(mid_std), fmt(mid_ref), "Staging 驗收 + 首批 20 張圖交付"),
    ("尾款", "20%", fmt(fin_std), fmt(fin_ref), "Production 上線 + 50 張圖全站接入 + 交付文件"),
    ("合計", "100%", fmt(LIST_TOTAL), fmt(REFERRAL_TOTAL), ""),
], [2, 1.5, 2.5, 2.5, 7])

h2(doc, "5.2 月度服務")
bullets(doc, [
    "每月 1 日前預付當月費用；首月可於簽約時與訂金一併或分開支付",
    "付款方式：銀行轉帳 / FPS；可提供香港商業發票（如適用）",
    "逾期：超過 14 日未付當期款項，承辦方可暫停所有服務直至清付",
])

# ── 六、不包含（另計）──
h1(doc, "六、不包含項目（另計）")
table(doc, ["項目", "參考"], [
    ("Meta / Google 廣告投放費", "客戶直接支付平台；代操含於月費"),
    ("Vercel Pro、網域年費", "各約 USD 20+/月、HKD 200–400/年"),
    ("醫師現場專業攝影", "另議（可取代 AI 診所圖）"),
    ("醫療廣告合規法律意見", "另議；文案合規責任見 FAQ"),
    ("CMS 後台 / 會員系統 / 24×7 駐場", "不包含"),
    ("超出合約範圍之開發或設計", f"HKD {HOURLY_DEV}/hr 或 HKD {HOURLY_IMAGE}/張"),
], [8, 8])

# ── 七、交付物 ──
h1(doc, "七、交付物")
bullets(doc, [
    "靜態網站原始碼（GitHub 私有倉庫存取權）",
    "Production 部署（oakvilles.com）及 DNS 設定指引",
    "sitemap.xml、robots.txt、Schema 設定",
    f"全站 {IMAGE_COUNT} 張 WebP/JPG 視覺資產",
    "設計指南、GA4 事件對照表、上線驗收清單",
    "月度服務：內容產出檔案、廣告月報、會議紀要（如適用）",
    "一次性項目：上線後 30 日缺陷保固（限原始 SOW 範圍內之程式錯誤，不含需求變更）",
])

# ── 八、條款摘要 ──
h1(doc, "八、條款摘要")
bullets(doc, [
    "本報價為技術開發及數碼行銷服務，不含醫療或法律意見；客戶對對外文案、圖片之合規性負最終責任",
    "AI 生成圖片不含可識別真實病患；醫師肖像以客戶提供素材為基礎",
    "第三方平台（Vercel、Meta、Google 等）政策變更，承辦方協助適配但不對平台決策負責",
    "需求變更超出 SOW 須事前書面確認及報價；未確認之工作不構成交付義務",
    "尾款及月費付清後，客戶享有交付物之約定使用權；開源套件依各自授權",
    "本報價 30 日內有效；轉介五折於簽約時確認",
    f"月度合約最少 {MIN_CONTRACT_MONTHS} 個月；期滿前 30 日書面通知可不續約",
    "爭議解決：雙方先協商；協商不成，以香港法律為準",
])

# ── 九、常見問題 FAQ（保障雙方，詳盡說明）──
h1(doc, "九、常見問題與條款說明（FAQ）")
para(doc, "以下條款旨在釐清雙方權責，避免日後爭議。簽約即視為同意。", size=10)

h2(doc, "9.1 付款與合約")
qa_block(doc, [
    (
        "若客戶延遲付款，承辦方有何權利？",
        f"訂金未付：承辦方不開始工作。中期或月費逾期超過 14 日：承辦方可暫停網站更新、行銷服務及廣告代操，"
        f"並保留追討權；因暫停導致之廣告空窗、排名波動或數據中斷，承辦方不承擔責任。"
        "尾款未付清：承辦方可暫緩移交完整源碼或管理權限，直至款項結清。",
    ),
    (
        "已付訂金後，客戶可取消項目嗎？退款如何計算？",
        "訂金為項目啟動及檔期保留，原則上不可退。若客戶主動取消，已產出之工時、素材及第三方成本按實際進度結算，"
        "餘額可退還；若進度已超過已收款項，客戶須補付差額。月度合約首月預付後，當月已開始之工作不予退還。",
    ),
    (
        "轉介五折點計？",
        f"經介紹之客戶，簽約時確認轉介來源，一次性項目享五折（HKD {fmt(REFERRAL_TOTAL)}）。月度服務不適用。",
    ),
    (
        "月度合約可否提早終止？",
        f"合約最少 {MIN_CONTRACT_MONTHS} 個月。若客戶於最低合約期內單方面終止，須付清剩餘合約期之月費，或按雙方書面協議之提前終止費用（以較高者為準）。"
        "承辦方因客戶嚴重違約（如長期欠款、要求違法文案）可即時終止，已付費用不予退還。",
    ),
])

h2(doc, "9.2 範圍、變更與驗收")
qa_block(doc, [
    (
        "什麼算「需求變更」？如何計費？",
        "超出本報價及 SOW 之新頁面、新功能、文案方向改動、額外語言版本、設計風格重做均屬變更。"
        f"變更須事前書面確認報價；開發按 HKD {HOURLY_DEV}/hr、圖片按 HKD {HOURLY_IMAGE}/張計。"
        "口頭或即時通訊之變更指示，在未書面確認前不構成承辦方義務。",
    ),
    (
        "驗收標準是什麼？客戶可以無限次修改嗎？",
        "網站以 SOW 功能清單及跨裝置 UAT 為驗收標準。視覺資產含 2 輪修訂；超出部分另計。"
        "客戶須於 Staging 驗收後 7 個工作天內以書面確認或一次列明修改清單；逾期視為該階段驗收通過。"
        "零散、反覆之修改若超出合理修訂範圍，承辦方可按工時另計。",
    ),
    (
        "若客戶遲遲不提供素材或反饋，工期如何計算？",
        "承辦方依客戶提供素材及確認之速度推進。因客戶延遲提供文案、圖片、帳號權限或審批而導致之工期順延，"
        "不視為承辦方違約；付款里程碑仍按合約執行，除非雙方書面同意調整。",
    ),
    (
        "上線後 30 日保固涵蓋什麼？",
        "僅涵蓋 SOW 範圍內之程式缺陷（如連結失效、預約表單錯誤、明顯排版崩壞），由承辦方免費修復。"
        "不含：新功能、內容更新、第三方平台故障、客戶或第三方自行修改程式碼導致之問題、"
        "瀏覽器外掛或客戶端環境造成之異常。保固期後維護依月度合約或按工時另計。",
    ),
])

h2(doc, "9.3 行銷、廣告與成效")
qa_block(doc, [
    (
        "月費是否保證預約量或廣告 ROAS？",
        "不保證。數碼行銷受競爭、季節、平台演算法及客戶預算影響。"
        "承辦方義務為按方案產出約定內容（4 IG + 4 FB + 2 Blog）、執行廣告代操及提供數據報告，"
        "而非保證特定預約數、排名或投資回報。客戶應合理預期並配合提供合規素材與審批。",
    ),
    (
        "廣告費用如何處理？最低預算建議？",
        "Meta / Google 廣告費由客戶以自身企業帳戶直接支付平台，發票及稅務由平台處理。"
        f"代操服務含於月費 HKD {fmt(MONTHLY_RETAINER)}，不含廣告 spend。"
        "建議試跑期各平台 HKD 3,000–8,000/月；實際預算由客戶決定，預算不足可能影響投放效果。",
    ),
    (
        "社交帳號及廣告帳戶歸誰擁有？",
        "Instagram、Facebook 專頁、Google Ads、GA4、Meta Business 等帳戶應登記於客戶名下。"
        "承辦方經客戶授權代為操作；合約終止時，承辦方須交還管理權限，不得扣留帳戶。"
        "承辦方自有之內容模板、內部工具及未交付之草稿不在移交範圍。",
    ),
    (
        "客戶要求違反醫療廣告規範的文案怎麼辦？",
        "承辦方有權拒絕製作或投放含「保證治癒」「誇大療效」等違規表述之素材。"
        "客戶堅持使用違規文案，承辦方可終止相關服務且不承擔因此導致之帳戶封禁或法律後果；已付月費不予退還。",
    ),
    (
        "每月 4+4 post 及 2 篇文章，若客戶未按時審批怎麼辦？",
        "承辦方按月度計劃提交草稿；客戶須於 5 個工作天內審批。"
        "若因客戶未審批導致無法發布，仍視為承辦方已履行該月產出義務（以交付草稿為準），除非雙方書面同意順延。",
    ),
])

h2(doc, "9.4 知識產權、資料與安全")
qa_block(doc, [
    (
        "網站源碼及圖片版權歸誰？",
        "尾款付清後，客戶享有專案交付之客製程式碼、頁面內容及已交付圖片於本項目範圍內之使用權。"
        "開源函式庫依各自授權；承辦方通用方法論、未交付之草稿及內部工具仍屬承辦方。"
        "客戶不可將交付物轉售為白標產品予第三方，除非另簽授權協議。",
    ),
    (
        "客戶自行修改網站後出問題，承辦方是否負責？",
        "否。客戶或第三方修改程式碼、主機設定或 DNS 後產生之故障，不在保固及月度維護範圍，"
        f"恢復工作按 HKD {HOURLY_DEV}/hr 另計。建議重大修改前先與承辦方確認。",
    ),
    (
        "客戶資料如何處理？",
        "承辦方僅在履行合約所需範圍內處理客戶提供之資料（文案、圖片、Analytics 等），"
        "不主動分享予第三方，合約終止後按客戶要求刪除或交還，法律另有規定除外。",
    ),
    (
        "若 Vercel、Meta、Google 等平台故障或改政策？",
        "承辦方合理協助適配，但不對平台中斷、帳戶封禁、政策變更導致之損失負責。"
        "因客戶違反平台政策（含廣告素材）導致之後果由客戶自行承擔。",
    ),
])

h2(doc, "9.5 其他")
qa_block(doc, [
    (
        "發票及報稅如何處理？",
        "承辦方可按香港法例提供商業發票（如適用）。客戶須自行處理其側之會計及稅務申報。"
        "第三方平台廣告費之發票由平台開立予客戶。",
    ),
    (
        "不可抗力（如疫情、天災）如何處理？",
        "因不可抗力導致無法履約，履約期限順延；若超過 60 日仍無法恢復，任何一方可書面終止，"
        "按已完成工作比例結算，互不追究違約責任（已產生之不可退成本除外）。",
    ),
    (
        "本報價與口頭承諾不一致時，以哪份為準？",
        "以本報價單及雙方簽署之正式合約 / SOW 為準；口頭或即時通訊之承諾，未經書面納入合約者不具約束力。",
    ),
])

# ── 十、確認簽署 ──
h1(doc, "十、確認簽署")
sig = doc.add_table(rows=5, cols=2)
sig.style = "Table Grid"
sig.rows[0].cells[0].text = "客戶 Client"
sig.rows[0].cells[1].text = "承辦方 Provider"
sig.rows[1].cells[0].text = "頤安本草 · 伍厚臻中醫師"
sig.rows[1].cells[1].text = "[承辦方名稱]"
sig.rows[2].cells[0].text = "授權代表\n\n_________________________"
sig.rows[2].cells[1].text = "授權代表\n\n_________________________"
sig.rows[3].cells[0].text = "日期\n\n_________________________"
sig.rows[3].cells[1].text = "日期\n\n_________________________"
sig.rows[4].cells[0].text = "簽署\n\n_________________________"
sig.rows[4].cells[1].text = "簽署\n\n_________________________"
for row in sig.rows:
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                style_run(r, 10)

doc.add_paragraph()
para(
    doc,
    "本文件為商務報價及條款說明，供雙方評估及簽約參考。金額按香港本地診所數碼項目市場水平制定。",
    size=8,
)

doc.save(OUTPUT)
shutil.copy2(OUTPUT, OUTPUT_DESKTOP)
os.makedirs(os.path.dirname(OUTPUT_PRESENTATIONS), exist_ok=True)
shutil.copy2(OUTPUT, OUTPUT_PRESENTATIONS)
print(f"Saved: {OUTPUT}")
print(f"Copied: {OUTPUT_DESKTOP}")
print(f"Copied: {OUTPUT_PRESENTATIONS}")
