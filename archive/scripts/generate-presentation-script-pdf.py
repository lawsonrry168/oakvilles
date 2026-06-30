# -*- coding: utf-8 -*-
"""Convert Oakville-Integrated-Presentation-Script.md to styled PDF."""

import os
import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = Path(r"c:\Users\Steriod\Desktop\oakvilles")
MD = BASE / "presentations" / "Oakville-Integrated-Presentation-Script.md"
DOCX_OUT = BASE / "presentations" / "Oakville-Integrated-Presentation-Script.docx"
PDF_OUT = BASE / "presentations" / "Oakville-Integrated-Presentation-Script.pdf"
PDF_DESKTOP = Path(r"c:\Users\Steriod\Desktop\新增資料夾 (2)") / "Oakville-Integrated-Presentation-Script.pdf"


def style_run(run, size=11, bold=False, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_rich_paragraph(doc, text, size=11, bold=False, italic=False, bullet=False):
    """Parse **bold** and `code` in a line."""
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
    else:
        p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25

    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            style_run(r, size, True, italic=italic)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            style_run(r, size - 0.5, False, RGBColor(0x2A, 0x46, 0x3C), italic)
        else:
            r = p.add_run(part)
            style_run(r, size, bold, italic=italic)


def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            break
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows


def build_docx():
    text = MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:])
            style_run(r, 18, True, RGBColor(0x2A, 0x46, 0x3C))
            p.paragraph_format.space_after = Pt(8)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:])
            style_run(r, 14, True, RGBColor(0x2A, 0x46, 0x3C))
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:])
            style_run(r, 12, True)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith("> "):
            add_rich_paragraph(doc, line[2:], size=10, italic=True)
        elif line.startswith("- "):
            add_rich_paragraph(doc, line[2:], bullet=True)
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(line, style="List Number")
            for r in p.runs:
                style_run(r, 11)
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            i -= 1
            rows = parse_table(table_lines)
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                t.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        t.rows[ri].cells[ci].text = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell)
                        for p in t.rows[ri].cells[ci].paragraphs:
                            for r in p.runs:
                                style_run(r, 9, ri == 0)
                doc.add_paragraph()
        elif line.strip() in ("---", ""):
            pass
        elif line.startswith("*") and line.endswith("*"):
            add_rich_paragraph(doc, line.strip("*"), size=9, italic=True)
        else:
            add_rich_paragraph(doc, line)

        i += 1

    doc.save(DOCX_OUT)
    return DOCX_OUT


def build_pdf(docx_path):
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(PDF_OUT))
        return "docx2pdf"
    except Exception as e1:
        pass

    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs(str(PDF_OUT.resolve()), FileFormat=17)
        doc.Close()
        word.Quit()
        return "win32com"
    except Exception as e2:
        raise RuntimeError(f"PDF conversion failed. Install MS Word or: pip install docx2pdf\n{e2}")


if __name__ == "__main__":
    docx = build_docx()
    method = build_pdf(docx)
    PDF_DESKTOP.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(PDF_OUT, PDF_DESKTOP)
    print(f"DOCX: {DOCX_OUT}")
    print(f"PDF:  {PDF_OUT}")
    print(f"Copy: {PDF_DESKTOP}")
    print(f"Method: {method}")
