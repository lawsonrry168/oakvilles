# -*- coding: utf-8 -*-
"""Generate website rebuild proposal for 頤安本草."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os
import shutil

OUTPUT = os.path.join(r"c:\Users\Steriod\Desktop\oakvilles", "Oakville-Website-Proposal.docx")
OUTPUT_DESKTOP = r"c:\Users\Steriod\Desktop\Oakville-Website-Proposal.docx"
today = datetime.date(2026, 6, 20)


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def style_run(run, size=10.5, bold=False):
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    style_run(r, 20, True)
    r.font.color.rgb = RGBColor(0x2A, 0x46, 0x3C)


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 14, True)
    r.font.color.rgb = RGBColor(0x2A, 0x46, 0x3C)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 11.5, True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_p(doc, text, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            style_run(r, 10)


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        set_cell_shading(t.rows[0].cells[i], "D5E8F0")
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                style_run(r, 9, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    style_run(r, 9)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

add_title(doc, "頤安本草 · 伍厚臻中醫師\n官方網站重建方案書")
add_p(doc, "Website Rebuild Proposal", size=11)
add_p(doc, f"日期：{today.year} 年 {today.month} 月 {today.day} 日　｜　客戶：頤安本草 · 伍厚臻中醫師　｜　網域：oakvilles.com", size=9)
doc.add_paragraph()

# 1
add_h1(doc, "一、背景與目的")
add_p(
    doc,
    "現行網站（oakvilles.com）為單頁式結構，資訊深度不足，難以承接「症狀搜尋」及「醫師名字搜尋」兩類主要流量。"
    "本方案以同業標竿養康中醫館（honscmc.com）為參考，梳理其資訊架構與獲客邏輯，"
    "並依頤安本草「中環高端、一人一方、循本溯源」之定位，重建一套可擴展、可量度、合規表述的官網系統。"
)
add_p(doc, "本文件目的：供內部決策及對外溝通，說明舊站不足、對標分析、新站改善重點及後續建議。", size=10)

# 2
add_h1(doc, "二、對標網站分析：養康中醫館（honscmc.com）")
add_h2(doc, "2.1 網站設計與資訊架構")
add_bullets(doc, [
    "平台：WordPress，頁面數量多，導航分為「專業療程」「診療範圍」「醫師療法」「中醫冷知識」四層下拉，覆蓋暗瘡、濕疹、減肥、豐胸、痛症、婦科等長尾關鍵字。",
    "視覺：以綠色臨床感為主，Hero 使用診所實景輪播，首屏即列出 WhatsApp 預約與營業時間，轉化路徑短。",
    "首屏訊息：「尖沙咀中醫診所 · 暗瘡濕疹皮膚專科 · 中醫美顏針 · 減肥豐胸」—— 地區 + 高需求科別一次講清。",
    "信任元件：13 年營運、年診量 8,000+、中文大學學歷、GMP 中藥、首次皮膚科免診金等，分散於首頁、FAQ 及專科頁。",
    "內容深度：各專科均有獨立著陸頁；首頁 FAQ 逾 10 題，利於 Google 摘要；「中醫冷知識」Blog 持續更新。",
    "第三方工具：Joinchat WhatsApp 小工具、Facebook／YouTube 連結、「用家留言」專頁嵌入 IG 貼文。",
])

add_h2(doc, "2.2 行銷與獲客策略（重點）")
add_p(
    doc,
    "養康的行銷並非只靠官網 SEO，而是「官網承接 + 社交種草 + 優惠降低首診門檻」的組合：",
    size=10,
)
add_table(
    doc,
    ["策略", "具體做法", "目的"],
    [
        ("SEO 長尾佈局", "每個症狀／療法獨立頁面，title 含地區 + 症狀 + 服務", "截取「暗瘡 中醫 尖沙咀」等搜尋流量"),
        ("WhatsApp 優先", "24 小時 WhatsApp、Joinchat 浮窗、全站「立即預約」", "符合香港用戶習慣，縮短決策路徑"),
        ("首診優惠", "皮膚科首次免診金", "降低試診成本，提高廣告轉化率"),
        ("IG／KOL 種草", "邀請 lifestyle／美容類 KOL 到店體驗，以「第一次針灸」「拔火罐體驗」等素人化口吻發文，標記 @hons_cmc", "製造口碑感，非硬銷廣告"),
        ("部落客／論壇", "PIXNET、Blogspot 等長文體驗文，含優惠碼與預約連結", "搜尋引擎補充外鏈與品牌詞佔位"),
        ("付費廣告著陸", "特定療程（如 AI 艾灸）獨立銷售頁，配合 Meta／Google 廣告", "流量導向高毛利療程"),
        ("文案外包", "聘請專業 copywriter 撰寫療程頁（公開 portfolio 可查）", "統一轉化語言與合規表述"),
    ],
    [3.5, 7.5, 5.5],
)
add_p(
    doc,
    "「素人化」IG 策略觀察：官網「用家留言」頁收錄多位 IG 用戶（如 lifestyle 類帳號）的到店體驗貼文，"
    "內容以第一人稱分享針灸、拔罐、艾灸過程，語氣接近真實用家而非官方廣告，"
    "再導回官網或 WhatsApp。此舉同時服務社交證明與 IG 探索流量。",
    size=9.5,
)
add_p(
    doc,
    "參考來源：honscmc.com 官網結構；用家留言頁；公開部落客體驗文（如 PIXNET heimen、Blogspot aminn613）；"
    "copywriter portfolio（georgeown.com/portfolio/hscm/）。",
    size=8.5,
)

add_h2(doc, "2.3 可借鑑之處（設計面）")
add_bullets(doc, [
    "以「症狀」而非「科別抽象名詞」作為 SEO 入口。",
    "FAQ 結構化，覆蓋地點、資歷、收費、預約方式。",
    "全站一致之 WhatsApp 預約動線。",
    "Blog 作為持續內容與內部連結樞紐。",
    "首屏即展示營業時間、地點、專科範圍。",
])

add_h2(doc, "2.4 不適合直接照搬之處（定位差異）")
add_bullets(doc, [
    "養康偏重美容塑形（美顏針、豐胸、減肥）及促銷語言；頤安本草定位「中環高端、循本溯源」，應維持克制、專業語氣。",
    "養康網站資訊密度高但視覺較擠；新站採東方禪意美學，強調私密、高端、留白。",
    "養康部分 efficacy 表述較進取；新站須符合香港中醫廣告合規，避免「數週見效」等絕對承諾。",
])

# 3
add_h1(doc, "三、現行網站（oakvilles.com）不足分析")
add_table(
    doc,
    ["維度", "現況", "影響"],
    [
        ("資訊架構", "實質為單頁網站，導航僅 5 項（首頁、關於、專科、專欄、地圖）", "無法覆蓋「頸痛 中醫 中環」等長尾搜尋"),
        ("專科覆蓋", "首頁僅展示皮膚、生育兩張卡片；標語卻提及痛症", "訊息不一致，痛症／內科流量無著陸頁"),
        ("SEO 技術", "無 sitemap、結構化資料、症狀專頁；OG 設定有限", "搜尋引擎難以理解網站結構與本地 relevance"),
        ("轉化設計", "僅 WhatsApp 浮動按鈕，無分步預約、無收費透明", "用戶需自行組織問題，流失率高"),
        ("信任建立", "缺 Google 評價展示、診所環境圖、流程說明、FAQ", "首次到訪者難以完成「5 秒信任測試」"),
        ("內容品質", "文案有筆誤（如「暗瘌」）；含「數週內可見皮損消退」", "專業感與合規風險"),
        ("品牌識別", "以「伍厚臻註冊中醫師」為主，未建立「頤安本草」品牌系統", "與高端診所定位不符，難以記憶"),
        ("聯絡資訊", "列出 28818182／28818280，與現行 WhatsApp 6734 9532 策略不一致", "多渠道混亂，不利追蹤"),
        ("量度能力", "無 GA4 事件、無 CTA 追蹤", "無法評估哪個入口帶來預約"),
        ("行動版", "有基本響應式，但 Hero 輪播為主，CTA 層級不明", "手機搜尋流量佔比高，轉化效率偏低"),
    ],
    [2.8, 5.5, 6.2],
)

# 4
add_h1(doc, "四、新站重建方案摘要")
add_p(
    doc,
    "新站（本地重建版本）已按上述差距完成架構與核心功能，定位為「中環高端中醫診所官網 + 可量度預約系統」，"
    "共 30 頁靜態網站，技術棧為 HTML5、自訂 CSS 設計系統（dongfang.css）、原生 JavaScript。",
    size=10,
)

add_h2(doc, "4.1 資訊架構（對標養康之深度，保留頤安之本位）")
add_table(
    doc,
    ["模組", "頁面", "對標借鑑"],
    [
        ("首頁", "Hero、四專科、診所 gallery、評價、計算器、預約", "首屏信任 + 多 CTA（養康）"),
        ("診症專科", "痛症、皮膚、婦科、內科 + 4 療法頁", "專科著陸頁（養康「診療範圍」）"),
        ("症狀專頁", "濕疹、暗瘡、失眠、備孕、頸痛、坐骨神經痛等 7 頁", "長尾 SEO（養康各症狀頁）"),
        ("內容", "Blog 4 篇、最新消息、FAQ", "中醫冷知識 + FAQ（養康）"),
        ("信任", "關於醫師、診所環境、流程收費、Google 評價", "資歷 + 環境 + 收費透明"),
        ("轉化", "兩步驟預約 funnel、WhatsApp、sticky CTA", "WhatsApp 優先（養康）"),
    ],
    [2.5, 7.5, 5.5],
)

add_h2(doc, "4.2 設計定位差異")
add_table(
    doc,
    ["項目", "養康中醫館", "頤安本草（新站）"],
    [
        ("視覺", "綠色臨床、促銷感", "paper／pine／cinnabar 東方禪意"),
        ("受眾", "尖沙咀、美容減肥、年輕女性", "中環上班族、高端、全科調理"),
        ("語氣", "促銷、療效導向", "循本溯源、一人一方、合規克制"),
        ("品牌", "養康中醫館", "頤安本草 · 伍厚臻中醫師"),
    ],
    [3, 5.5, 7],
)

# 5
add_h1(doc, "五、新站較舊站之具體改善")
add_table(
    doc,
    ["類別", "舊站", "新站"],
    [
        ("頁面規模", "1 頁為主", "30 頁，含 7 症狀著陸頁"),
        ("品牌", "個人醫師名義", "頤安本草品牌系統 + 印章識別"),
        ("SEO", "基本 title", "sitemap、robots、JSON-LD、canonical、OG"),
        ("本地搜尋", "無「中環中醫」專頁", "central-hk.html 等本地頁"),
        ("預約", "WhatsApp 連結", "2 步 funnel + 計算器預填 + sessionStorage 串接"),
        ("收費", "未展示", "明碼價目 + 即時估算計算器"),
        ("評價", "無", "Google 5.0 三則評價 + schema Review"),
        ("量度", "無", "GA4 dataLayer 事件（cta_click、booking_submit 等）"),
        ("內容", "2 專科卡片", "4 專科 + 7 症狀 + Blog + FAQ"),
        ("合規", "有效能承諾用語", "設計指南明確禁止誇大表述"),
        ("視覺", "通用 SPA 風", "自訂設計系統 + 50 張品牌圖規劃"),
        ("行動版", "基本響應式", "Mobile-first + sticky 底部 CTA"),
    ],
    [2.5, 5.5, 7.5],
)

# 6
add_h1(doc, "六、行銷策略建議（參考養康，適配頤安）")
add_p(doc, "以下為新站上線後之建議方向，可納入月度行銷服務（HKD 10,000／月）執行：", size=10)

add_h2(doc, "6.1 建議採納")
add_bullets(doc, [
    "Google Business Profile 優化：中環定位、服務項目、評價邀請。",
    "WhatsApp 為主轉化通道，所有 CTA 統一追蹤。",
    "FAQ + 症狀頁持續補充，目標截取「[症狀] 中醫 中環」搜尋。",
    "Blog 每月 1–2 篇，題材對應濕疹、暗瘡、失眠、備孕等核心科別。",
    "IG @oakville.wellness 與官網互導，展示診所日常及養生知識（非硬銷）。",
])

add_h2(doc, "6.2 可選討論（素人化／KOL）")
add_p(
    doc,
    "養康模式：邀請 micro-KOL 以「體驗分享」形式發布 IG 貼文，官網「患者好評」或專欄引用（需取得同意）。"
    "頤安若採用，建議：",
    size=10,
)
add_bullets(doc, [
    "選擇與「中環上班族／皮膚困擾／備孕」相關之垂直 KOL，避免純美容促銷調性。",
    "內容聚焦「診症體驗」「醫師解說」「診所環境」，避免疗效承诺。",
    "所有合作貼文標記贊助或廣告（符合 HK 廣告規範）。",
    "官網只展示已授權內容，與 Google 真實評價分開呈現。",
])

add_h2(doc, "6.3 不建議照搬")
add_bullets(doc, [
    "首診免診金等高折扣促銷（與高端定位不符，除非短期獲客實驗）。",
    "豐胸、減肥等高敏感療程作主要 SEO 入口（與伍醫師現有專長重心不同）。",
    "大量 efficacy 導向之論壇軟文，合規風險較高。",
])

# 7
add_h1(doc, "七、交付與後續")
add_table(
    doc,
    ["階段", "內容", "狀態"],
    [
        ("第一階段", "30 頁網站、SEO 基礎、預約 funnel、GA4 就緒", "已開發"),
        ("第二階段", "50 張品牌視覺資產製作及全站接入", "規劃完成，待製作"),
        ("第三階段", "Production 上線 oakvilles.com", "待部署"),
        ("持續", "月度行銷策劃 + 網站維護（HKD 10,000／月）", "可選合約"),
    ],
    [2.5, 8, 3],
)

add_h1(doc, "八、結論")
add_p(
    doc,
    "養康中醫館的優勢在於「內容厚度 + WhatsApp 轉化 + 社交種草」三者結合；"
    "現行 oakvilles.com 則內容與技術基礎不足，未能承接搜尋與轉介流量。"
    "新站已在資訊架構、SEO、轉化量度及品牌視覺上系統性補足差距，"
    "並刻意保持中環高端、合規、專業之定位，與養康之促銷密集型策略區隔。"
    "建議下一步：完成視覺資產接入、正式上線，並按月度方案持續優化本地 SEO 與內容更新。",
    size=10.5,
)

doc.add_paragraph()
add_p(doc, "— 完 —", size=9)
add_p(
    doc,
    "附：對標網站 https://honscmc.com/　｜　現行網站 https://oakvilles.com/　｜　新站預覽（本地重建）",
    size=8,
)

doc.save(OUTPUT)
shutil.copy2(OUTPUT, OUTPUT_DESKTOP)
print(f"Saved: {OUTPUT}")
