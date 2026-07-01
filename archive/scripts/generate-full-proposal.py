# -*- coding: utf-8 -*-
"""Generate strategy + website + marketing proposal (no pricing; plain language)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os
import shutil

BASE = r"c:\Users\Steriod\Desktop\oakvilles"
OUT = os.path.join(BASE, "Oakville-Full-Proposal.docx")
OUT_PRES = os.path.join(BASE, "presentations", "Oakville-Full-Proposal.docx")

IMAGE_PLANNED = 50
IMAGE_CURRENT = 60
today = datetime.date(2026, 7, 1)
PINE = RGBColor(0x2A, 0x46, 0x3C)


def set_cell_shading(cell, fill: str):
    s = OxmlElement("w:shd")
    s.set(qn("w:fill"), fill)
    s.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(s)


def style_run(run, size=10.5, bold=False, color=None):
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    style_run(r, 20, True, PINE)


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 14, True, PINE)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 11.5, True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def para(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size, bold)
    p.paragraph_format.space_after = Pt(5)


def bullets(doc, items, size=10):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            style_run(r, size)


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        set_cell_shading(t.rows[0].cells[i], "D5E8F0")
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                style_run(r, 9, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    style_run(r, 9)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

# ── Cover ──
title(doc, "頤安本草 · 伍厚臻中醫師\n官網重建方案書")
para(doc, "策略 · 網站 · 行銷", size=12)
para(
    doc,
    f"{today.year} 年 {today.month} 月　｜"
    "頤安本草（Oakville Wellness）　｜"
    "oakvilles.com",
    size=9,
)
para(
    doc,
    "本文件說明重建方向與具體安排；費用及合約條款另載於《綜合報價單》，兩份文件分開閱讀。",
    size=9,
)
doc.add_paragraph()

h2(doc, "目錄")
bullets(doc, [
    "第一部分　策略 — 為何重建、對標借鑑、品牌定位",
    "第二部分　網站 — 內容架構、設計風格、預約體驗、現況",
    "第三部分　行銷 — 每月節奏、廣告配合、成效衡量、路線圖",
    "結語",
], size=9.5)

# ═══════════════════════════════════════
# PART I — 策略
# ═══════════════════════════════════════
h1(doc, "第一部分　策略")

h2(doc, "1.1 執行摘要")
para(
    doc,
    "現行官網內容太少，客人從 Google 搜尋「中環 濕疹 中醫」或朋友轉介來時，"
    "往往找不到足夠資訊，也難以放心預約。",
)
para(
    doc,
    "本方案為頤安本草重建一套完整官網：共 62 個中英頁面，"
    "清楚介紹專科、常見症狀、診所環境與收費，"
    "並把預約統一導向 WhatsApp 6734 9532。",
    True,
)
para(
    doc,
    "整體定位：中環高端、專業合規、循本溯源——"
    "學習同業在內容與預約上的優點，但不走促銷、折扣路線。",
)

h2(doc, "1.2 我們要解決的問題")
table(doc, ["現況", "對診所的影響"], [
    ("網站幾乎只有一頁", "搜尋「症狀 + 中環 + 中醫」時難以被看見"),
    ("沒有各症狀專頁", "濕疹、暗瘡、失眠、備孕等客人無專門入口"),
    ("預約只靠一個浮動按鈕", "客人不清楚流程與收費，容易離開"),
    ("缺少評價、環境圖、常見問題", "首次到訪者難以建立信任"),
    ("品牌以個人名義為主", "「頤安本草」難以被記住"),
    ("無法知道哪個渠道帶來預約", "日後投放廣告時無從優化"),
], [6, 10])

h2(doc, "1.3 行業參考：養康中醫館")
para(doc, "養康官網在業內表現突出，以下做法值得參考，但不必照搬：", size=10)
table(doc, ["值得學習", "頤安不跟隨"], [
    ("每種常見症狀都有獨立介紹頁", "美容塑形、豐胸減肥作主打"),
    ("常見問題寫清楚地點、收費、預約", "首診免診金、高折扣促銷"),
    ("全站用 WhatsApp 預約，路徑一致", "誇大療效、論壇軟文灌水"),
    ("持續更新專欄文章", "資訊過密、促銷感太強的版面"),
    ("社交媒體與官網互相導流", ""),
], [6, 10])

h2(doc, "1.4 三方簡要對照")
table(doc, ["", "養康（標竿）", "舊官網", "新官網"], [
    ("定位", "尖沙咀·皮膚美容·促銷", "個人名義", "中環高端·頤安本草"),
    ("視覺", "綠色臨床·密度高", "通用單頁", "霧藍宣紙·墨藍·留白"),
    ("頁面", "多頁、內容厚", "約 1 頁", "62 頁（中+英）"),
    ("搜尋", "症狀專頁齊全", "幾乎沒有", "8 個症狀專頁 + 中環專頁"),
    ("預約", "全站 WhatsApp", "只有浮動鈕", "兩步預約 + 全站一致"),
    ("信任", "資歷、診量、FAQ", "較薄弱", "Google 評價、環境圖、明碼收費"),
], [2, 3.5, 3.5, 7])

h2(doc, "1.5 品牌定位")
table(doc, ["", "養康", "頤安本草"], [
    ("視覺感受", "綠色、臨床、資訊多", "霧藍宣紙、墨藍文字、留白、高端克制"),
    ("主要客人", "尖沙咀、美容需求", "中環上班族、重視專業與私隱"),
    ("說話方式", "促銷、見效導向", "循本溯源、一人一方、合規克制"),
    ("品牌名稱", "養康中醫館", "頤安本草 · 伍厚臻中醫師"),
], [3, 5.5, 7.5])

# ═══════════════════════════════════════
# PART II — 網站
# ═══════════════════════════════════════
h1(doc, "第二部分　網站")

h2(doc, "2.1 新官網是什麼")
para(
    doc,
    "一套為頤安本草度身訂做的診所官網：手機與電腦均可流暢瀏覽，"
    "中文為主、並設完整英文版，方便外籍客人。"
    "網站已建置完成並上線預覽，待正式網域 oakvilles.com 切換後即可全面使用。",
)

h2(doc, "2.2 網站裡有什麼")
table(doc, ["板塊", "客人能看到什麼"], [
    ("首頁", "診所介紹、四大專科、診所照片、客人評價、診金試算"),
    ("專科與療法", "痛症、皮膚、婦科、內科；針灸、中藥、艾灸、拔罐等"),
    ("常見症狀", "濕疹、暗瘡、失眠、備孕、頸痛、坐骨神經痛等專頁"),
    ("中環專頁", "交通、地點、為何選擇中環門診"),
    ("醫師與診所", "伍醫師資歷、診所實景、就診流程與收費"),
    ("專欄與消息", "養生文章、診所公告"),
    ("常見問題", "針灸痛不痛、如何服藥、如何預約等"),
    ("英文版", "以上內容的英文對應頁面"),
], [3, 13])

h2(doc, "2.3 預約與收費體驗")
bullets(doc, [
    "全站預約統一：WhatsApp +852 6734 9532，並附標準問候語，方便客人一鍵開啟對話",
    "兩步預約：先選科別與日期，再填聯絡方式，系統自動組好訊息供客人發送",
    "診金計算器：客人可自行試算，並一鍵帶入預約",
    "手機底部固定預約欄，瀏覽任何頁面都可隨時聯絡",
    "收費透明：流程與價目頁清楚列出，減少查詢往返",
])

h2(doc, "2.4 搜尋與分享")
para(doc, "技術細節已處理妥當，客人無需關心；對診所的實際好處如下：", size=10)
bullets(doc, [
    "Google 可完整收錄各頁，有利「症狀 + 地區」搜尋",
    "分享至 WhatsApp、Facebook 時會顯示體面預覽圖與簡介",
    "中英文頁面互相對應，有利本地與外籍搜尋",
    "已備妥數據統計基礎，日後可知道多少人點擊預約（待接上統計帳號）",
])

h2(doc, "2.5 信任如何建立")
bullets(doc, [
    "展示 Google 5.0 評價（附合規免責說明）",
    "診所環境相簿、就診流程圖",
    "伍醫師學歷與臨床經驗",
    "常見問題解答地點、資歷、收費、預約",
    "文案遵守合規：不承諾「根治」「數週見效」等表述",
])

h2(doc, "2.6 設計風格與色調")
para(
    doc,
    "整體定調為「湖水晨霧宣紙」的東方雅韻：偏冷、偏霧藍灰的紙感背景，"
    "搭配深墨藍文字與適度留白，營造中環高端、寧靜、專業的診所印象——"
    "刻意與養康的綠色臨床促銷感區隔。",
)
h2(doc, "色彩運用")
table(doc, ["用途", "色調描述", "參考色"], [
    ("頁面背景", "淡霧藍白宣紙，帶微藍漸層與細緻紙紋", "#EBF3F7"),
    ("標題與正文", "深墨藍灰，沉穩易讀", "#2A3A44"),
    ("深色區塊／預約按鈕", "墨藍（非鮮豔促銷綠）", "#1A2832"),
    ("眉題與副標", "霧藍灰點綴", "#6E8494"),
    ("連結與強調", "低飽和藍灰", "#5A7A8C"),
    ("首頁品牌浮水印", "香檳金褐（中英 Hero 背景字）", "#C5B390"),
], [3.5, 7.5, 5])

h2(doc, "字體與識別")
bullets(doc, [
    "中文正文：Noto Serif TC 宋體風格，字級偏大，方便各年齡層閱讀",
    "中文品牌／印章：Oakville 隸書字體 + PNG 印章圖（頤安本草、候診醫師等主題印）",
    "英文品牌副標：Cormorant Garamond 襯線斜體（OAKVILLE WELLNESS）",
    "英文正文：Noto Serif，與中文版氣質一致",
    "版面：手機優先；導覽、底部預約欄固定，方便隨時聯絡",
])

h2(doc, "2.7 視覺圖片")
table(doc, ["項目", "說明"], [
    ("規劃總量", f"約 {IMAGE_PLANNED} 張品牌圖片（診所、專科、專欄、流程等）"),
    ("現已上線", f"約 {IMAGE_CURRENT} 張主圖（部分為佔位，待換成正式攝影）"),
    ("圖片風格", "與霧藍宣紙底調和；避免高飽和促銷色塊"),
    ("下一步", "以真實診所與醫師照片，逐步替換佔位圖"),
], [4, 12])

h2(doc, "2.8 現況：已完成與待辦")
table(doc, ["狀態", "內容"], [
    ("已完成", "62 頁中英網站、預約動線、搜尋基礎、評價與收費頁、雲端上線預覽"),
    ("待完成", "正式網域 oakvilles.com 指向新站"),
    ("待完成", "接上網站數據統計，以便衡量廣告成效"),
    ("持續進行", "品牌圖片批次更新、內容微調"),
], [2.5, 13.5])

h2(doc, "2.9 舊官網 → 新官網（一覽）")
table(doc, ["項目", "舊官網", "新官網"], [
    ("頁面數量", "約 1 頁", "62 頁（中+英）"),
    ("品牌", "個人醫師", "頤安本草品牌系統"),
    ("視覺", "通用模板、輪播為主", "霧藍宣紙 + 隸書印章 + 墨藍系"),
    ("症狀搜尋", "無", "8 個症狀專頁"),
    ("預約", "單一連結", "兩步預約 + 全站一致"),
    ("收費", "未列出", "明碼價 + 試算"),
    ("評價", "無", "Google 5.0"),
    ("常見問題", "無", "獨立 FAQ 頁"),
    ("手機體驗", "基本", "底部預約欄、版面優化"),
], [2.5, 5, 8.5])

# ═══════════════════════════════════════
# PART III — 行銷
# ═══════════════════════════════════════
h1(doc, "第三部分　行銷")

h2(doc, "3.1 整體思路")
para(
    doc,
    "用社交貼文種草、用 Google 與 Facebook／Instagram 廣告找準客人，"
    "把人帶到官網對應的症狀或專科頁，最後統一經 WhatsApp 6734 9532 預約。",
    True,
)
para(doc, "成功與否，主要看有多少人點擊 WhatsApp 開始預約對話——其餘數據用來優化內容與廣告。")

h2(doc, "3.2 每月固定做什麼")
table(doc, ["工作", "頻率", "說明"], [
    ("Instagram 貼文", "每週 1 則", "每月共 4 則"),
    ("Facebook 貼文", "每週 1 則", "每月共 4 則（可與 IG 共用素材）"),
    ("官網專欄文章", "每兩週 1 篇", "每月共 2 篇，配合搜尋與專業形象"),
    ("Facebook 廣告", "持續優化", "活動設定、受眾、素材測試、每月報告"),
    ("Instagram 廣告", "持續優化", "同上，含 Reels／限時動態"),
    ("Google 搜尋廣告", "持續優化", "「中環 + 症狀 + 中醫」等高意向關鍵字"),
    ("Google 商家檔案", "每月維護", "地圖曝光、評價、營業資訊"),
    ("數據檢討", "每月 1 次", "約 60 分鐘線上會議，檢視成效與下月計劃"),
], [3.5, 2.5, 10])

h2(doc, "3.3 每月四週節奏")
bullets(doc, [
    "第 1 週：發布症狀科普（例如濕疹與體質的關係）",
    "第 2 週：發布診所信任內容（環境、資歷、評價）+ 官網第 1 篇文章",
    "第 3 週：季節養生主題 + 加強表現最好的廣告",
    "第 4 週：教客人如何 WhatsApp 預約 + 官網第 2 篇文章 + 本月總結",
    "每則內容結尾提醒：WhatsApp 6734 9532 查詢檔期",
])

h2(doc, "3.4 廣告與官網如何配合")
table(doc, ["想吸引的客人", "廣告點進去看", "平台建議"], [
    ("濕疹、暗瘡困擾", "對應症狀專頁", "Facebook、Instagram、Google"),
    ("備孕、婦科調理", "備孕專頁", "Facebook、Instagram、Google"),
    ("中環找中醫", "中環專頁", "Google 搜尋為主"),
    ("外籍人士", "英文首頁或皮膚專科", "Google、Instagram"),
], [4, 5, 7])

para(
    doc,
    "廣告費用由診所直接支付予 Facebook、Instagram、Google；"
    "代為設定與優化廣告已包含在月度服務內（詳見報價單）。"
    "試跑期建議 Facebook、Instagram、Google 各保留基本預算，"
    "跑順後按預約成本調高或調低。",
    size=10,
)

h2(doc, "3.5 我們不做什麼")
bullets(doc, [
    "免診金、高折扣促銷（與高端定位不符）",
    "豐胸、減肥等作主要推廣入口",
    "論壇軟文、誇大療效表述",
    "保證特定預約人數或廣告回報（受市場與季節影響，只能盡力優化）",
])

h2(doc, "3.6 上線後三個月路線圖")
table(doc, ["時段", "重點"], [
    ("第 1–2 週", "正式網域上線、接上網站統計、優化 Google 地圖商家檔案"),
    ("第 3–6 週", "Facebook、Instagram、Google 各開一組廣告試跑；發首兩篇專欄"),
    ("第 2–3 月", "保留有效廣告、調整預算；每月固定內容與數據檢討"),
], [3, 13])

# ═══════════════════════════════════════
# 結語
# ═══════════════════════════════════════
h1(doc, "結語")
para(
    doc,
    "舊官網內容單薄，難以承接搜尋與轉介；"
    "新官網已在架構、內容、預約動線與視覺呈現（霧藍宣紙、墨藍文字、隸書印章）上系統補足，"
    "並刻意保持中環高端、合規、專業的調性。",
    size=10.5,
)
para(
    doc,
    "建議下一步：完成正式網域切換、接上數據統計，"
    "並按月度方案持續發布內容與優化廣告，"
    "讓「找到我們」的客人，能順利「預約到診」。",
    size=10.5,
)

doc.add_paragraph()
para(
    doc,
    "— 完 —\n"
    "費用及合約：請參閱《Oakville-Combined-Quotation》綜合報價單。\n"
    "新站預覽：oakvilles.vercel.app　｜　WhatsApp +852 6734 9532",
    size=8,
)

doc.save(OUT)
os.makedirs(os.path.dirname(OUT_PRES), exist_ok=True)
shutil.copy2(OUT, OUT_PRES)
print(f"Saved: {OUT}")
print(f"Copied: {OUT_PRES}")
