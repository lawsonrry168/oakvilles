# -*- coding: utf-8 -*-
"""Generate 6-month IG / FB / Blog content calendar DOCX — full blog articles."""

import datetime
import importlib.util
import os
import shutil
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"c:\Users\Steriod\Desktop\oakvilles"
OUT = os.path.join(BASE, "Oakville-Content-Calendar-6M.docx")
OUT_PRES = os.path.join(BASE, "presentations", "Oakville-Content-Calendar-6M.docx")
TODAY = datetime.date(2026, 6, 24)

GLOBAL_STYLE = (
    "GLOBAL STYLE — Oakville Wellness 頤安本草: Premium Hong Kong Central TCM clinic aesthetic. "
    "Warm cream paper tone (#F7F2E7), deep pine green (#2A463C), subtle cinnabar red (#A23A2E), "
    "muted ochre gold (#AE8A47). Zen-like, private, calm, high-end editorial — NOT generic hospital stock. "
    "Soft natural window light, gentle shadows, shallow depth of field. Negative space for typography overlay. "
    "Subtle rice-paper texture in highlights. No harsh flash, no watermarks, NO TEXT IN IMAGE. "
    "Warm desaturated greens, ivory highlights. Mood: 禪意、私密、專業、中環高端. Photorealistic, 8K detail."
)

EXISTING_AVOID = [
    "blog/eczema-from-within.html — 濕疹為何要從內調",
    "blog/acne-face-zones.html — 暗瘡位置與臟腑",
    "blog/insomnia-tcm-guide.html — 失眠中醫分型",
    "blog/fertility-stress.html — 壓力型難孕",
]


def _load_data():
    path = os.path.join(BASE, "archive", "scripts", "content-calendar-data.py")
    spec = importlib.util.spec_from_file_location("content_calendar_data", path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def style_run(run, size=10.5, bold=False, color=None):
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    style_run(r, 20, True, RGBColor(0x2A, 0x46, 0x3C))


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 14, True, RGBColor(0x2A, 0x46, 0x3C))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 12, True, RGBColor(0x2A, 0x46, 0x3C))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 11, True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)


def para(doc, text, size=10, bold=False):
    for block in text.split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(block)
        style_run(r, size, bold)
        p.paragraph_format.space_after = Pt(3)


def prompt_block(doc, scene: str, ratio: str):
    full = f"{scene}\n\n{GLOBAL_STYLE}\n\nAspect ratio: {ratio}."
    p = doc.add_paragraph()
    r = p.add_run(full)
    style_run(r, 9)
    p.paragraph_format.space_after = Pt(8)


def render_full_article(doc, blog: dict):
    """Render complete blog body from sections."""
    para(doc, f"分類：{blog.get('category', '養生專欄')}", 9)
    para(doc, f"建議 slug：{blog['slug']}", 9)
    para(doc, f"著陸頁：{blog['landing']}", 9)
    doc.add_paragraph()

    for sec in blog["sections"]:
        h2(doc, sec["h2"])
        for ptext in sec["paragraphs"]:
            para(doc, ptext, 10)

    h2(doc, "結語")
    para(doc, blog.get("closing", ""), 10)


def overview_table(doc, months):
    headers = ["月份", "IG", "FB", "Blog", "季節主題"]
    rows = []
    for m in months:
        blogs = sum(1 for w in m["weeks"] if w.get("blog"))
        rows.append((m["label"], "4 則", "4 則", f"{blogs} 篇", m["season"]))
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        set_cell_shading(t.rows[0].cells[i], "D5E8F0")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def blog_index_table(doc, months):
    h2(doc, "養生專欄索引（全新主題）")
    headers = ["月份", "週次", "標題", "slug"]
    blogs = []
    for m in months:
        for w in m["weeks"]:
            if w.get("blog"):
                b = w["blog"]
                blogs.append((m["label"], f"W{w['w']}", b["title"], b["slug"]))
    t = doc.add_table(rows=1 + len(blogs), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        set_cell_shading(t.rows[0].cells[i], "D5E8F0")
    for ri, row in enumerate(blogs):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def build():
    data = _load_data()
    months = data.MONTHS
    hashtags_ig = data.HASHTAGS_IG
    hashtags_fb = data.HASHTAGS_FB

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    title(doc, "頤安本草 · 半年圖文內容規範（文案場景對齊版）")
    para(doc, "Instagram · Facebook · 官網養生專欄 — 2026 年 7–12 月", 11, True)
    para(
        doc,
        f"制定日期：{TODAY.year} 年 {TODAY.month} 月 {TODAY.day} 日　｜"
        "全部主題與現有官網專欄及上一版日曆不重複",
        9,
    )
    doc.add_paragraph()

    h1(doc, "與現有內容區隔")
    para(doc, "以下官網已發布專欄，本日曆刻意迴避相同角度與標題：", 10)
    for item in EXISTING_AVOID:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            style_run(r, 10)

    h1(doc, "使用說明")
    for b in [
        "圖片：每則 IG/FB Prompt 含「場景對應文案」句 + 視覺描述，須與當週 caption 一致。",
        "社交文案：IG / FB 各 24 則，主題已更新（見逐週條目）。",
        "養生專欄：12 篇完整正文（可直接貼上官網或再潤飾），非僅大綱。",
        "節奏：每月 4 IG + 4 FB；Blog 於第 2、4 週發布。",
        f"統一 CTA：{data.CTA}",
    ]:
        p = doc.add_paragraph(b, style="List Bullet")
        for r in p.runs:
            style_run(r, 10)

    h1(doc, "三種渠道規格")
    para(doc, "Instagram：1:1（1080×1080）· Hero 吸睛構圖 · 主體佔畫面 55–65% · 預留標題留白", 10)
    para(doc, "Facebook：4:5（1080×1350）· 直式封面感 · 主體偏上、下方留白給 CTA", 10)
    para(doc, "養生專欄：Hero 16:10 · 全文約 900–1200 字 · 含 FAQ 與結語", 10)

    h1(doc, "半年總覽")
    overview_table(doc, months)
    blog_index_table(doc, months)

    for m in months:
        doc.add_page_break()
        h1(doc, m["label"] + "　" + m["season"])

        for w in m["weeks"]:
            h2(doc, f"第 {w['w']} 週 · {w['pillar']} · {w['topic']}")
            para(doc, f"建議著陸頁：{w['landing']}", 9)

            h3(doc, "A. Instagram")
            para(doc, "【圖片 Prompt · 1:1 · 場景須對應上方文案】", 9, True)
            ig_scene = data.SOCIAL_IG_LEAD + w.get("img_copy_hook", "") + w["img_scene"]
            prompt_block(doc, ig_scene, "1:1 square, 1080×1080")
            para(doc, "【文案】", 9, True)
            para(doc, w["ig"], 10)
            para(doc, f"Hashtags：{hashtags_ig}", 8)

            h3(doc, "B. Facebook")
            para(doc, "【圖片 Prompt · 4:5 · 場景須對應上方文案】", 9, True)
            fb_scene = data.SOCIAL_FB_LEAD + w.get("img_copy_hook", "") + w["img_scene"]
            prompt_block(doc, fb_scene, "4:5 portrait, 1080×1350")
            para(doc, "【文案】", 9, True)
            para(doc, w["fb"], 10)
            para(doc, f"Hashtags：{hashtags_fb}", 8)

            if w.get("blog"):
                b = w["blog"]
                h3(doc, "C. 官網養生專欄（完整正文）")
                para(doc, f"標題：{b['title']}", 11, True)
                para(doc, "【圖片 Prompt · 16:10 Hero】", 9, True)
                prompt_block(doc, b["img_scene"], "16:10, 1600×1000")
                render_full_article(doc, b)

    h1(doc, "附錄：發布前合規檢查")
    for item in [
        "無「治癒」「根治」「保證見效」「100%」",
        "療效表述：調理、支援、改善；加「效果因人而異」",
        "圖片無內嵌文字、無皮膚病灶特寫、無他人診所 Logo",
        "醫師資歷：註冊中醫 003769",
        "與現有四篇官網專欄標題不重複",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            style_run(r, 10)

    doc.save(OUT)
    os.makedirs(os.path.dirname(OUT_PRES), exist_ok=True)
    shutil.copy2(OUT, OUT_PRES)

    blog_count = sum(1 for m in months for w in m["weeks"] if w.get("blog"))
    print(f"Saved: {OUT}")
    print(f"Copied: {OUT_PRES}")
    print(f"IG/FB: {len(months) * 4} each, Blogs (full text): {blog_count}")


if __name__ == "__main__":
    build()
