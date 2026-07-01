# -*- coding: utf-8 -*-
"""Merge Strategy Pack + Digital Marketing Plan into one simplified DOCX."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os
import shutil

BASE = r"c:\Users\Steriod\Desktop\oakvilles"
OUT = os.path.join(BASE, "Oakville-Marketing-Summary.docx")
OUT_PRES = os.path.join(BASE, "presentations", "Oakville-Marketing-Summary.docx")
today = datetime.date(2026, 7, 1)
MONTHLY_RETAINER = 8_000
MONTHLY_RETAINER_PREV = 10_000
META_FB = 1_000
META_IG = 1_000
GOOGLE_AD = 3_000


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def style_run(run, size=10.5, bold=False, color=None):
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    style_run(r, 20, True, RGBColor(0x2A, 0x46, 0x3C))


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 14, True, RGBColor(0x2A, 0x46, 0x3C))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 11.5, True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def para(doc, text, size=10.5, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size, bold)
    p.paragraph_format.space_after = Pt(5)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            style_run(r, 10)


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        set_cell_shading(t.rows[0].cells[i], "D5E8F0")
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                style_run(r, 9, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    style_run(r, 9)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    title(doc, "頤安本草 · 行銷執行摘要")
    para(
        doc,
        "合併《策略文件包》與《網絡行銷方案》之精簡版 — 一頁看懂目標、節奏、廣告與成效",
        11,
    )
    para(
        doc,
        f"日期：{today.year} 年 {today.month} 月 {today.day} 日　｜"
        "oakvilles.com　｜"
        "WhatsApp 6734 9532　｜"
        "@oakville.wellness",
        9,
    )
    doc.add_paragraph()

    # ── 1. One-page overview ──
    h1(doc, "一、一句話說清楚")
    para(
        doc,
        "在中環建立「伍厚臻中醫師 · 高端、合規、可信」形象，透過社交種草、"
        "官網 SEO 與 Meta／Google 廣告，把人導到官網，最後以 WhatsApp 6734 9532 預約。",
        11,
        True,
    )
    table(doc, ["項目", "內容"], [
        ("目標客群", "中環／金鐘／半山上班族 + 外籍人士（/en/）"),
        ("主打症狀", "濕疹、暗瘡、失眠、備孕、頸肩痛"),
        ("唯一 KPI", "whatsapp_click — 用戶點擊 WhatsApp 預約連結"),
        ("品牌語氣", "循本溯源、一人一方、合規克制（不用「根治」「免診金」）"),
    ], [3.5, 12.5])

    h2(doc, "我們做 / 不做")
    table(doc, ["做", "不做"], [
        ("每月 4 IG + 4 FB + 2 篇官網文章", "免診金、論壇軟文、豐胸減肥 SEO"),
        ("症狀頁承接廣告與搜尋", "與高端定位不符的促銷語"),
        ("Meta + Google 以 WA 轉化為目標", "只看瀏覽量、不追蹤來源"),
    ], [7, 9])

    # ── 2. Monthly rhythm ──
    h1(doc, "二、每月固定節奏")
    table(doc, ["週次", "社交貼文主題", "官網", "廣告營運"], [
        ("W1", "症狀科普（例：濕疹與體質）", "—", "檢視上週報表"),
        ("W2", "診所信任（環境、資歷、Google 5.0）", "發第 1 篇 Blog", "調整受眾／著陸頁"),
        ("W3", "季節養生（祛濕、補腎等）", "—", "加碼表現最佳廣告組"),
        ("W4", "預約教學（2 步 WhatsApp 預約）", "發第 2 篇 Blog", "月度 ROI 檢討"),
    ], [1.2, 5.5, 3.5, 4.3])
    bullets(doc, [
        "IG 與 FB 可共用同一套圖文／Reels，結尾統一 CTA：WhatsApp 6734 9532",
        "Blog 每篇 800–1200 字，連結對應症狀頁（/conditions/）與專科頁",
        "每月更新 Google 商家檔案 1–2 則；表現最佳貼文可餵給 Meta 廣告",
    ])

    # ── 3. Ads simplified ──
    h1(doc, "三、付費廣告（精簡版）")
    table(doc, ["平台", "做什麼", "著陸到哪裡"], [
        ("Meta", "症狀受眾轉化 + 再營銷曾訪官網者", "濕疹／暗瘡／備孕等 /conditions/"),
        ("Google 搜尋", "「中環 + 症狀 + 中醫」高意向關鍵字", "對應症狀頁或 central-hk"),
        ("Google 商家", "本地曝光、評價、每週動態", "Google Maps / 地圖頁"),
    ], [2.5, 6.5, 7.5])
    para(doc, "受眾重點：香港 25–55 歲、中環一帶、對中醫／養生／皮膚／備孕有興趣。", 10)
    para(doc, "追蹤：網站已埋 6 個事件；開廣告前須完成 GA4 + GTM + Meta Pixel，以 whatsapp_click 為轉化。", 10)

    # ── 4. Measurement ──
    h1(doc, "四、成效怎麼看")
    table(doc, ["指標", "意思", "怎樣算好（參考）"], [
        ("whatsapp_click", "點了 WhatsApp — 核心 KPI", "月增 15–25%"),
        ("booking_submit", "填了 2 步預約表單", "約 WA 點擊的 20–30%"),
        ("症狀頁停留", "內容是否吸引對的人", "> 1 分鐘"),
        ("廣告 CPA", "每次 WA 點擊花多少", "試跑 4–6 週後訂目標"),
    ], [3.5, 6.5, 6.5])
    bullets(doc, [
        "P0：在 js/site-config.js 填入 GA4 ID，並用 GTM 接上 Pixel",
        "每月看：WA 點擊總數、來源（社交／廣告／自然）、哪個著陸頁最好",
    ])

    # ── 5. Budget ──
    h1(doc, "五、預算參考（每月）")
    para(
        doc,
        f"月度行銷及維護服務費 HKD {MONTHLY_RETAINER:,}（原 {MONTHLY_RETAINER_PREV:,}）；"
        "內容製作（8 則社交 + 2 文）已含於月費。",
        10,
        True,
    )
    h2(doc, "月費配置（服務費）")
    table(doc, ["項目", "HKD", "說明"], [
        ("Meta Facebook 代操", f"{META_FB:,}", "活動、受眾、素材、Pixel、月報"),
        ("Meta Instagram 代操", f"{META_IG:,}", "Reels/Stories、IG 受眾優化"),
        ("Google 搜尋代操", f"{GOOGLE_AD:,}", "品牌+中環+症狀 Search"),
        ("內容製作（8 則社交 + 2 文）", "月費包含", "不再另收 3,000–6,000"),
        ("策略·維護·會議", "月費包含", "GBP、GA4 月報、網站小修、月度檢討"),
        ("月度服務費合計", f"{MONTHLY_RETAINER:,}", "不含平台廣告投放費"),
    ], [5.5, 2.5, 8])
    h2(doc, "建議廣告投放（客戶直付平台）")
    table(doc, ["項目", "HKD"], [
        ("Meta Facebook", f"{META_FB:,}"),
        ("Meta Instagram", f"{META_IG:,}"),
        ("Google 搜尋", f"{GOOGLE_AD:,}"),
        ("內容製作", "—（已含月費）"),
        ("客戶每月合計（服務+廣告）", f"約 {MONTHLY_RETAINER + META_FB + META_IG + GOOGLE_AD:,}"),
    ], [10, 6])
    h2(doc, "價格調整前後")
    table(doc, ["項目", "調整前", "調整後"], [
        ("月度服務費", f"{MONTHLY_RETAINER_PREV:,}", f"{MONTHLY_RETAINER:,}"),
        ("內容製作", "另計 3,000–6,000", "月費包含"),
        ("Meta", "合併 5,000–8,000", f"FB {META_FB:,} + IG {META_IG:,}"),
        ("Google", "5,000–8,000", f"{GOOGLE_AD:,}"),
    ], [5, 5, 6])
    para(doc, "建議試跑 4–6 週，依 whatsapp_click 成本再調整廣告 spend。", 10)

    # ── 6. Competitor one-pager ──
    h1(doc, "六、相對養康（精簡對照）")
    table(doc, ["我們較強", "養康較強 → 我們怎麼應對"], [
        ("中環高端 + 雙語官網 62 頁", "KOL 多 → 中階 KOL 季度合作 + 自有 4+4 post"),
        ("合規品牌、GA4 量度就緒", "首診優惠 → 不跟價；強調專業與體驗"),
        ("預約 funnel、診金計算器", "內容更厚 → 每月 2 篇 Blog 持續累積"),
    ], [6.5, 9.5])
    para(
        doc,
        "較易成交：中環上班族、重視合規與專業、症狀搜尋高意向。"
        "較難成交：只要最便宜、一定要尖沙咀、美容塑形需求 — 非我們主戰場。",
        10,
    )

    # ── 7. 90-day roadmap ──
    h1(doc, "七、90 天路線圖")
    table(doc, ["階段", "時間", "重點"], [
        ("P0 打底", "第 1–2 週", "GA4 · GTM · Pixel · Search Console · Google 商家優化"),
        ("試跑", "第 3–6 週", "Meta + Google 各 1 活動；首 2 篇 Blog；內容日曆跑順"),
        ("優化", "第 7–10 週", "加碼最佳廣告；社交最佳帖加推"),
        ("穩定", "第 11–12 週", "月度 ROI 檢討；調整預算與下月題材"),
    ], [2, 2.5, 11.5])

    # ── 8. P0 checklist ──
    h1(doc, "八、立即行動清單（P0）")
    bullets(doc, [
        "填入 GA4 Measurement ID，驗證 whatsapp_click 有記錄",
        "Search Console 提交 sitemap；優化 Google 商家檔案",
        "定稿首月 8 則社交 + 2 篇 Blog 標題",
        "Meta / Google 各開 1 個轉化活動，著陸濕疹或備孕頁試跑",
        "每則內容結尾統一：WhatsApp 6734 9532 查詢檔期",
    ])

    para(
        doc,
        "— 完 —\n"
        "詳盡版請參：Oakville-Digital-Marketing-Plan.docx、Oakville-Strategy-Pack.docx",
        9,
    )

    doc.save(OUT)
    os.makedirs(os.path.dirname(OUT_PRES), exist_ok=True)
    shutil.copy2(OUT, OUT_PRES)
    print(f"Saved: {OUT}")
    print(f"Copied: {OUT_PRES}")


if __name__ == "__main__":
    build()
