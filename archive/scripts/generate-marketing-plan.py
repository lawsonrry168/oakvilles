# -*- coding: utf-8 -*-
"""Generate digital marketing plan DOCX for 頤安本草."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os
import shutil

OUTPUT = os.path.join(r"c:\Users\Steriod\Desktop\oakvilles", "Oakville-Digital-Marketing-Plan.docx")
OUTPUT_DESKTOP = r"c:\Users\Steriod\Desktop\Oakville-Digital-Marketing-Plan.docx"
OUTPUT_PRESENTATIONS = os.path.join(
    r"c:\Users\Steriod\Desktop\oakvilles\presentations", "Oakville-Digital-Marketing-Plan.docx"
)
today = datetime.date(2026, 6, 20)


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def style_run(run, size=10.5, bold=False):
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    style_run(r, 20, True)
    r.font.color.rgb = RGBColor(0x2A, 0x46, 0x3C)


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 14, True)
    r.font.color.rgb = RGBColor(0x2A, 0x46, 0x3C)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, 11.5, True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_p(doc, text, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            style_run(r, 10)


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        set_cell_shading(t.rows[0].cells[i], "D5E8F0")
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                style_run(r, 9, True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    style_run(r, 9)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

add_title(doc, "頤安本草 · 網絡行銷方案")
add_p(doc, "Digital Marketing Plan — Oakville Wellness", size=11)
add_p(
    doc,
    f"日期：{today.year} 年 {today.month} 月 {today.day} 日　｜"
    "官網：oakvilles.com　｜"
    "IG：@oakville.wellness　｜"
    "FB：facebook.com/oakvillewellness",
    size=9,
)
doc.add_paragraph()

# 一
add_h1(doc, "一、方案目標")
add_table(doc, ["層次", "目標"], [
    ("品牌", "建立「中環高端、循本溯源、合規專業」形象"),
    ("流量", "症狀搜尋 + 地區搜尋 + 社交種草導向官網"),
    ("轉化", "以 WhatsApp 6734 9532 預約為唯一核心 KPI"),
    ("長期", "官網 SEO 資產累積（症狀頁 + 專欄文章）"),
], [3, 13])
add_p(doc, "不採用：免診金促銷、論壇軟文、豐胸減肥等與品牌定位不符的獲客手段。", size=10)

# 二
add_h1(doc, "二、每月內容節奏（固定產出）")
add_h2(doc, "2.1 產出總量")
add_table(doc, ["渠道", "頻率", "每月"], [
    ("Instagram", "每週 1 篇", "4 post"),
    ("Facebook", "每週 1 篇", "4 post"),
    ("官網養生專欄", "每兩週 1 篇", "2 篇文章"),
], [4, 4, 4])
add_p(
    doc,
    "IG 與 FB 可共用同一套素材（圖文／輪播／Reels 擇一），依平台微調文案長度與 CTA 位置。",
    size=10,
)

add_h2(doc, "2.2 社交內容四大支柱（每月各 1 篇，輪替）")
add_table(doc, ["週次", "支柱", "範例主題", "著陸頁"], [
    ("W1", "症狀科普", "濕疹反覆？中醫從體質看根因", "/conditions/eczema/"),
    ("W2", "診所日常／信任", "診所環境、伍醫師資歷、Google 5.0 評價", "/about/ 或 /process/"),
    ("W3", "季節養生", "春夏祛濕、秋冬補腎、節氣調理", "當月 Blog 新文"),
    ("W4", "轉化導向", "如何預約、2 步 WhatsApp 預約示範", "首頁 / 或 /contact/"),
], [1.5, 3, 5.5, 4.5])
add_h2(doc, "2.3 文案原則")
add_bullets(doc, [
    "合規克制：用「調理」「改善」「支援」等表述，避免「根治」「100% 有效」",
    "每篇結尾固定 CTA：「WhatsApp 6734 9532 查詢檔期」+ 連結官網",
    "圖片風格：paper/pine 品牌色、診所實景、圖文卡片（非促銷感）",
])

add_h2(doc, "2.4 官網每月 2 篇文章")
add_p(
    doc,
    "定位：SEO 長尾 + 社交轉發素材 + 症狀頁內部連結。",
    size=10,
)
add_bullets(doc, [
    "題材優先：濕疹、暗瘡、失眠、備孕（深化）→ 頸痛、坐骨神經痛 → 中環上班族議題",
    "英文版 /en/blog/ 可隔月同步 1 篇（外籍受眾）",
    "每篇 800–1200 字，含 FAQ 小節",
    "內部連結至對應 /conditions/ 頁 + 相關 /services/ 專科",
    "文末 CTA：WhatsApp 預約 + 診金計算器；OG 圖 1200×630",
])

# 三
add_h1(doc, "三、Meta 廣告（Facebook + Instagram）")
add_h2(doc, "3.1 投放結構")
add_table(doc, ["活動類型", "目的", "著陸頁"], [
    ("轉化（主力）", "WhatsApp 點擊", "症狀頁（濕疹／暗瘡／備孕／失眠）"),
    ("流量（輔助）", "官網瀏覽 + 再營銷", "/blog/ 新文或首頁"),
    ("本地觸及", "中環 3–5 km 受眾", "/conditions/central-hk/"),
], [3.5, 4.5, 6.5])

add_h2(doc, "3.2 受眾設定")
add_bullets(doc, [
    "核心：香港 · 25–55 歲 · 中環／金鐘／半山／西環",
    "興趣：中醫、養生、健康、皮膚護理、備孕、痛症",
    "Lookalike：官網 WhatsApp 點擊者（Pixel 累積後）",
    "再營銷：7／14／30 天曾訪官網但未點 WA 的訪客",
])

add_h2(doc, "3.3 素材與追蹤")
add_bullets(doc, [
    "沿用每月 4 組社交 post（最佳表現者加預算）",
    "單圖／輪播／15–30 秒 Reels（診所環境、預約流程）",
    "文案 A/B：「中環 · 伍厚臻中醫師」vs「症狀關鍵字 + 合規描述」",
    "Meta Pixel 經 GTM 部署；自訂轉換 whatsapp_click",
    "建議上線 Conversion API（CAPI），提升 iOS 追蹤準確度",
])

# 四
add_h1(doc, "四、Google 廣告")
add_h2(doc, "4.1 投放結構")
add_table(doc, ["類型", "關鍵字方向", "著陸頁"], [
    ("Search · 品牌", "頤安、伍厚臻、Oakville 中醫", "首頁"),
    ("Search · 地區", "中環中醫、中環針灸、Central TCM", "/conditions/central-hk/"),
    ("Search · 症狀", "中環濕疹中醫、香港暗瘡中醫、備孕中醫", "對應 /conditions/"),
    ("Performance Max", "轉化優化（WhatsApp）", "多症狀頁 + 首頁"),
], [3.5, 5.5, 5.5])

add_h2(doc, "4.2 關鍵字與本地 SEO")
add_bullets(doc, [
    "高意向：「中環 + 症狀 + 中醫／針灸」",
    "排除：減肥、豐胸、免費、最便宜等低質流量",
    "英文（可選）：eczema TCM Central HK → /en/conditions/eczema/",
    "Google Business Profile：地址、營業時間、服務項目、定期發帖",
    "引導 Google 5.0 評價客戶留評，與廣告信任背書互補",
])

add_h2(doc, "4.3 追蹤設定")
add_bullets(doc, [
    "填入 SITE_GA4_ID（js/site-config.js）",
    "GA4 轉化事件：whatsapp_click、booking_submit",
    "Search Console 提交 sitemap",
    "GA4 轉化匯入 Google Ads → Maximize Conversions",
])

# 五
add_h1(doc, "五、渠道分工")
add_p(
    doc,
    "社交（IG/FB）→ 種草、信任、日常觸及　｜"
    "官網 Blog → SEO 長尾、深度內容、內部連結　｜"
    "症狀頁 → 廣告著陸、搜尋承接　｜"
    "Meta 廣告 → 症狀受眾 + 再營銷 + Reels　｜"
    "Google 搜尋 → 高意向「正在找中醫」　｜"
    "WhatsApp 6734 9532 → 唯一轉化終點",
    size=10,
)

# 六
add_h1(doc, "六、每月工作流")
add_table(doc, ["週", "社交", "官網", "廣告"], [
    ("W1", "IG+FB 症狀科普 post", "—", "檢查 Meta / Google 報表"),
    ("W2", "IG+FB 信任／診所 post", "發布第 1 篇 Blog", "調整著陸頁與受眾"),
    ("W3", "IG+FB 季節養生 post", "—", "加碼表現最佳 ad set"),
    ("W4", "IG+FB 預約 CTA post", "發布第 2 篇 Blog", "月度 ROI 檢討"),
], [1.2, 4.5, 4, 4])
add_h2(doc, "每月固定事項")
add_bullets(doc, [
    "檢視 whatsapp_click 數量與來源（hero / sticky / wa-float / 廣告）",
    "更新 Google Business Profile 1–2 則動態",
    "社交 best post 餵給 Meta 廣告素材庫",
])

# 七
add_h1(doc, "七、量度指標（KPI）")
add_table(doc, ["指標", "目標方向"], [
    ("whatsapp_click", "主要 KPI（社交 + 廣告 + 自然）"),
    ("booking_submit", "高意向次要轉化"),
    ("官網 Sessions", "月增（SEO + 廣告）"),
    ("症狀頁停留時間", "> 1 分鐘"),
    ("廣告 CPA", "單次 WA 點擊成本（依預算設定）"),
    ("社交互動率", "IG > 2%、FB > 1%（參考值）"),
], [4.5, 11])
add_p(
    doc,
    "官網已就緒 6 個 dataLayer 事件；GA4 ID 與 GTM（GA4 + Meta Pixel）為廣告開跑前 P0。",
    size=10,
)

# 八
add_h1(doc, "八、預算建議（參考）")
add_table(doc, ["項目", "建議月費（HKD）"], [
    ("Meta 廣告", "5,000 – 8,000"),
    ("Google Search", "5,000 – 8,000"),
    ("內容製作（4+4 post + 2 文）", "3,000 – 6,000（視外包／自製）"),
    ("合計", "約 13,000 – 22,000／月"),
], [8, 7.5])
add_p(doc, "可先以 Meta 3,000 + Google 3,000 試跑 4–6 週，依 WA 轉化成本再調配。", size=10)

# 九
add_h1(doc, "九、合規提醒（香港醫療廣告）")
add_bullets(doc, [
    "不得宣稱「治癒」「保證見效」",
    "中醫師資歷（003769）可展示，療效個案需匿名且措辭謹慎",
    "社交與廣告素材需與官網合規語氣一致",
    "英文受眾內容同步遵守相同原則",
])

# 十
add_h1(doc, "十、立即行動（P0）")
add_bullets(doc, [
    "GA4 + GTM + Meta Pixel 部署並驗證 whatsapp_click",
    "Search Console + Google Business Profile 提交與優化",
    "制定首月內容日曆（4 IG + 4 FB + 2 Blog 標題定稿）",
    "Meta / Google 各開 1 個轉化活動，著陸濕疹或備孕頁試跑",
    "每篇社交 post 與 Blog 文末統一導向 WhatsApp 6734 9532",
])

doc.save(OUTPUT)
shutil.copy2(OUTPUT, OUTPUT_DESKTOP)
os.makedirs(os.path.dirname(OUTPUT_PRESENTATIONS), exist_ok=True)
shutil.copy2(OUTPUT, OUTPUT_PRESENTATIONS)
print(f"Saved: {OUTPUT}")
print(f"Copied: {OUTPUT_DESKTOP}")
print(f"Copied: {OUTPUT_PRESENTATIONS}")
